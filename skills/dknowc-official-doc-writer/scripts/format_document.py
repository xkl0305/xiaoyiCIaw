#!/usr/bin/env python3
"""
公文排版脚本 v2.0.0
严格按照《广东省政务服务和数据管理局公文格式样本（试行）》进行自动排版

【v2.0.0 更新说明】
- 删除所有红头文件相关代码（红头+尾表由 template_generator.py 处理）
- 只保留普通格式排版核心逻辑
- 架构：template_generator.py 调用本文件的函数进行正文排版
"""

import argparse
import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path

import docx
import docx.opc.constants
from docx.enum.text import WD_BREAK
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_hyperlink(paragraph, url, text, font_name=None, font_size=None):
    """在段落中添加可点击的超链接"""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if font_name or font_size:
        rFonts = OxmlElement('w:rFonts')
        if font_name:
            rFonts.set(qn('w:eastAsia'), font_name)
            rFonts.set(qn('w:ascii'), font_name)
            rFonts.set(qn('w:hAnsi'), font_name)
        rPr.append(rFonts)
    if font_size:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(font_size * 2)))
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(int(font_size * 2)))
        rPr.append(szCs)

    # 超链接颜色（蓝色）
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    # 下划线
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink


def sanitize_url(raw_url):
    """清理 URL 后缀说明，避免把括号内说明写进超链接。"""
    url = raw_url.strip().rstrip('，。、；;,.')
    url = re.split(r'[（(]', url, maxsplit=1)[0]
    return url.strip().rstrip('，。、；;,.')


def is_internal_search_endpoint(url):
    """识别深知搜索接口地址，避免误当成知识专库链接输出。"""
    return "open.dknowc.cn/dependable/search" in url

DEFAULT_ORG = "广东省政务服务和数据管理局"
DEFAULT_DOC_PREFIX = "粤政数"
OUTPUT_DIR = os.path.expanduser("~/.openclaw/data/official-docs/output")
CONFIG_PATH = os.path.join(os.path.dirname(__file__), '..', 'config', 'format.json')
AI_DISCLAIMER_TEXT = "【AI生成提示】内容由AI生成，内容仅供参考。"


def load_format_config():
    """加载格式配置文件"""
    try:
        with open(CONFIG_PATH, 'r', encoding='utf-8') as f:
            return json.load(f)
    except FileNotFoundError:
        print(f"警告: 配置文件不存在 {CONFIG_PATH}，使用默认值")
        return None
    except json.JSONDecodeError as e:
        print(f"警告: 配置文件格式错误 {CONFIG_PATH}: {e}，使用默认值")
        return None


# 加载配置（全局变量）
FORMAT_CONFIG = load_format_config()


def get_font_config(element_name):
    """从配置获取字体设置，返回 (font_name, font_size, bold)"""
    if not FORMAT_CONFIG or element_name not in FORMAT_CONFIG:
        # 默认值
        defaults = {
            'title': ('方正小标宋简体', 22, False),
            'body': ('仿宋_GB2312', 16, False),
            'heading1': ('黑体', 16, False),
            'heading2': ('楷体_GB2312', 16, False),
            'heading3': ('仿宋_GB2312', 16, True),
        }
        return defaults.get(element_name, ('仿宋_GB2312', 16, False))
    
    config = FORMAT_CONFIG[element_name]
    return (
        config.get('font', '仿宋_GB2312'),
        config.get('size_pt', 16),
        config.get('bold', False)
    )


def get_page_margin():
    """从配置获取页边距，返回 (top_mm, bottom_mm, left_mm, right_mm)"""
    if FORMAT_CONFIG and 'page' in FORMAT_CONFIG:
        page = FORMAT_CONFIG['page']
        # 支持四边分别设置（国标 GB/T 9704）
        if 'margin_top_mm' in page:
            return (
                page.get('margin_top_mm', 37),
                page.get('margin_bottom_mm', 33),
                page.get('margin_left_mm', 28),
                page.get('margin_right_mm', 26),
            )
        # 兼容旧配置：统一边距
        unified = page.get('margin_mm', 25)
        return (unified, unified, unified, unified)
    return (37, 33, 28, 26)


def get_page_size():
    """从配置获取纸张尺寸，默认 A4。"""
    if FORMAT_CONFIG and 'page' in FORMAT_CONFIG:
        page = FORMAT_CONFIG['page']
        return (
            page.get('width_mm', 210),
            page.get('height_mm', 297),
        )
    return (210, 297)


def get_latin_font_config():
    """获取全文数字和字母字体。"""
    if FORMAT_CONFIG:
        body = FORMAT_CONFIG.get('body', {})
        if body.get('latin_font'):
            return body['latin_font']
    return 'Times New Roman'

MAX_INPUT_LENGTH = 50000
MAX_INPUT_LINES = 2000
MAX_OUTPUT_LINES = 3000


def get_next_version(output_path: str) -> str:
    """
    获取下一个版本号的输出路径

    Args:
        output_path: 原始输出路径

    Returns:
        带版本号的输出路径（如果文件已存在）
    """
    output_path = os.path.expanduser(output_path)
    output_dir = Path(output_path).parent
    filename = Path(output_path).stem  # 文件名（无扩展名）
    ext = Path(output_path).suffix     # 扩展名

    # 如果文件不存在，直接返回原路径
    if not Path(output_path).exists():
        return output_path

    # 文件已存在，需要添加版本号
    # 查找所有同名文件（包含版本号）
    pattern = f"{filename}_v*.docx"
    existing_files = list(output_dir.glob(pattern))

    # 如果没有版本号文件，从v1开始
    if not existing_files:
        return str(output_dir / f"{filename}_v1{ext}")

    # 提取所有版本号
    versions = []
    for file in existing_files:
        # 提取版本号（如"关于XXX的通知_v1.docx" -> 1）
        try:
            version_str = file.stem.split('_v')[-1]
            version_num = int(version_str)
            versions.append(version_num)
        except:
            continue

    # 计算下一个版本号
    next_version = max(versions) + 1 if versions else 1

    return str(output_dir / f"{filename}_v{next_version}{ext}")


try:
    from docx import Document
    from docx.shared import Pt, Cm, Mm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("错误: 缺少 python-docx 库，请安装: pip install python-docx")
    sys.exit(1)


def set_run_font(run, font_name, font_size, bold=False, color=None, latin_font=None, latin=False):
    """设置字体"""
    latin_font = latin_font or get_latin_font_config()
    display_font = latin_font if latin else font_name
    run.font.name = display_font
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), display_font)
    rFonts.set(qn('w:hAnsi'), display_font)
    rFonts.set(qn('w:cs'), display_font)


def add_formatted_text(paragraph, text, font_name, font_size, bold=False, color=None):
    """添加文本并确保数字、字母使用 Times New Roman。"""
    latin_font = get_latin_font_config()
    for segment in re.findall(r'[A-Za-z0-9]+|[^A-Za-z0-9]+', text):
        run = paragraph.add_run(segment)
        set_run_font(
            run,
            font_name,
            font_size,
            bold=bold,
            color=color,
            latin_font=latin_font,
            latin=bool(re.fullmatch(r'[A-Za-z0-9]+', segment)),
        )
    return paragraph


def set_paragraph_format(para, first_line_indent=False, indent_chars=None, alignment=None, line_spacing=None):
    """设置段落格式
    
    Args:
        para: 段落对象
        first_line_indent: 是否首行缩进（已废弃，请使用indent_chars）
        indent_chars: 缩进字符数（如2表示2字符）
        alignment: 对齐方式
        line_spacing: 行距（磅）
    """
    if indent_chars:
        # 同时写入常规缩进，兼容只读取 first_line_indent 的校验器。
        para.paragraph_format.first_line_indent = Cm(0.37 * indent_chars)
        # 使用Word的字符单位缩进（更准确）
        # 通过底层XML设置 w:firstLineChars 属性
        # 值的单位是1/100字符，所以2字符 = 200
        pPr = para._p.get_or_add_pPr()
        ind = pPr.ind
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        ind.set(qn('w:firstLineChars'), str(int(indent_chars * 100)))
    elif first_line_indent:
        # 兼容旧配置
        indent_cm = FORMAT_CONFIG['body']['first_line_indent_cm'] if FORMAT_CONFIG else 0.74
        para.paragraph_format.first_line_indent = Cm(indent_cm)
    if alignment is not None:
        para.alignment = alignment
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    if line_spacing:
        para.paragraph_format.line_spacing = Pt(line_spacing)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY


def set_outline_level(para, level):
    """设置段落的大纲级别
    
    Args:
        para: 段落对象
        level: 大纲级别（0-9，0为最高级）
               0 = 1级（对应 # 标题）
               1 = 2级（对应 ## 一、XXX）
               2 = 3级（对应 ### （一）XXX）
               3 = 4级（对应 #### 1.XXX）
               9 = 正文文本（默认）
    """
    pPr = para._p.get_or_add_pPr()
    outlineLvl = OxmlElement('w:outlineLvl')
    outlineLvl.set(qn('w:val'), str(level))
    pPr.append(outlineLvl)


def add_center_page_number(section):
    """在页脚居中添加页码。"""
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    set_run_font(run, '宋体', 14)

    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'PAGE'
    fld_separate = OxmlElement('w:fldChar')
    fld_separate.set(qn('w:fldCharType'), 'separate')
    text = OxmlElement('w:t')
    text.text = '1'
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_separate)
    run._r.append(text)
    run._r.append(fld_end)


def add_ai_disclaimer(doc):
    """在文档最末尾添加 AI 生成提示。"""
    if doc.paragraphs and doc.paragraphs[-1].text.strip() == AI_DISCLAIMER_TEXT:
        return

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = Pt(18)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    add_formatted_text(para, AI_DISCLAIMER_TEXT, '仿宋_GB2312', 10.5, color=RGBColor(0x80, 0x80, 0x80))


def get_line_type(line):
    """
    根据Markdown结构判断行类型
    
    返回值:
    - 'title': # 标题（公文标题）
    - 'heading1': ## 一级标题
    - 'heading2': ### 二级标题
    - 'heading3': #### 三级标题
    - 'empty': 空行
    - 'text': 普通文本
    """
    stripped = line.strip()
    
    # 空行
    if not stripped:
        return 'empty'
    
    if stripped.startswith('#### '):
        return 'heading3'
    if stripped.startswith('### '):
        return 'heading2'
    if stripped.startswith('## '):
        return 'heading1'
    if stripped.startswith('# ') and not stripped.startswith('## '):
        return 'title'

    # 兼容模型未加 Markdown 标记但使用公文常见标题格式的情况。
    if re.match(r'^[一二三四五六七八九十]+、[^。；;]{1,40}$', stripped):
        return 'heading1'
    if re.match(r'^（[一二三四五六七八九十]+）[^。；;]{1,40}$', stripped):
        return 'heading2'
    if re.match(r'^\d+[\.．、][^。；;]{1,40}$', stripped):
        return 'heading3'
    
    return 'text'


def strip_markdown_heading(line):
    """去掉Markdown标题标记"""
    stripped = line.strip()
    if stripped.startswith('#### '):
        return stripped[5:]
    if stripped.startswith('### '):
        return stripped[4:]
    if stripped.startswith('## '):
        return stripped[3:]
    if stripped.startswith('# ') and not stripped.startswith('## '):
        return stripped[2:]
    return stripped


def has_markdown_title(lines):
    """判断正文是否已经包含 Markdown 公文标题。"""
    return any(get_line_type(line.strip()) == 'title' for line in lines)


def infer_title_from_output_path(output_path):
    """当正文遗漏 # 标题时，从输出文件名推断公文标题。"""
    if not output_path:
        return None
    stem = Path(os.path.expanduser(output_path)).stem.strip()
    if not stem:
        return None
    stem = re.sub(r'_v\d+$', '', stem)
    stem = re.sub(r'_红头$', '', stem)
    if re.match(r'^公文_\d{8}_\d{6}$', stem):
        return None
    if len(stem) < 4 or len(stem) > 80:
        return None
    return stem


def strip_markdown_bold(text):
    """去掉Markdown加粗标记 **text** 或 __text__"""
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'__([^_]+)__', r'\1', text)
    return text


def is_attachment_line(line):
    """判断是否为附件行"""
    return bool(re.match(r'^附件[:：]', line.strip()))


def is_attachment_continuation(line):
    """判断是否为附件清单续行，如 2.×××。"""
    return bool(re.match(r'^\d+[\.．、]', line.strip()))


def set_attachment_list_format(para, continuation=False, line_spacing=None):
    """设置正文末尾附件说明格式。"""
    pPr = para._p.get_or_add_pPr()
    ind = pPr.ind
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:start'), "0")
    ind.set(qn('w:startChars'), "0")
    ind.set(qn('w:end'), "0")
    ind.set(qn('w:endChars'), "0")
    if continuation:
        para.paragraph_format.left_indent = Cm(0)
        para.paragraph_format.first_line_indent = None
        ind.set(qn('w:firstLineChars'), "112")
    else:
        para.paragraph_format.first_line_indent = Cm(0.74)
        ind.set(qn('w:firstLineChars'), "200")
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    if line_spacing:
        para.paragraph_format.line_spacing = Pt(line_spacing)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY


def is_attachment_section_title(line):
    """判断附件正文首页标题。"""
    return line.strip() == "附件"


def document_has_visible_content(doc):
    """判断文档中是否已经有正文内容，用于附件页前分页。"""
    return any(paragraph.text.strip() for paragraph in doc.paragraphs)


def is_signing_entity(line):
    """判断是否为落款单位
    
    落款单位特征：
    - 长度小于30
    - 以机构名结尾（局、厅、委、办、中心等）
    - 不包含具体场所（会议室、办公室等）
    - 不以句号结尾
    """
    stripped = line.strip()
    
    if len(stripped) >= 30:
        return False
    
    # 排除以句号结尾的（不是落款单位）
    if stripped.endswith('。'):
        return False
    
    # 排除包含具体场所的（会议室、服务中心等）
    exclude_keywords = ['会议室', '服务中心', '办事大厅', '窗口']
    if any(kw in stripped for kw in exclude_keywords):
        return False
    
    # 必须以机构名结尾
    agency_suffixes = ['局', '厅', '委', '办', '办公室', '中心', '院', '会', '组委会', '协会', '站', '所', '部', '处', '司', '署', '公司', '集团', '单位']
    if not any(stripped.endswith(suffix) for suffix in agency_suffixes):
        return False
    
    return True


def is_date_line(line):
    """判断是否为公文落款日期。"""
    stripped = re.sub(r'\s+', '', line.strip())
    patterns = [
        r'^\d{4}年(?:\d{1,2}|【[^】]+】)月(?:\d{1,2}|【[^】]+】)日$',
        r'^【[^】]+】年【[^】]+】月【[^】]+】日$',
        r'^(?:\d{4}|×{4}|XXXX)年(?:\d{1,2}|×{1,2}|XX)月(?:\d{1,2}|×{1,2}|XX)日$',
    ]
    return any(re.match(pattern, stripped) for pattern in patterns)


def should_right_align_date(lines, idx):
    """仅将正文落款附近的日期右对齐，避免正文日期误判。"""
    if not is_date_line(lines[idx].strip()):
        return False

    prev_non_empty = ""
    for prev_idx in range(idx - 1, -1, -1):
        candidate = lines[prev_idx].strip()
        if candidate:
            prev_non_empty = candidate
            break

    if prev_non_empty and is_signing_entity(prev_non_empty):
        return True

    next_non_empty = ""
    for next_idx in range(idx + 1, min(len(lines), idx + 4)):
        candidate = lines[next_idx].strip()
        if candidate:
            next_non_empty = candidate
            break

    return next_non_empty in ("[分页符]", "【素材使用情况】", "【知识专库链接】")


def is_contact_info(line):
    """判断是否为联系人信息"""
    return bool(re.match(r'^(?:（?联系人|联系电话|联系人电话|联系方式)[:：]', line.strip()))


def normalize_content_text(content_text):
    """规范输入换行，降低命令行传参导致整篇文本变成一段的风险。"""
    content_text = content_text.replace('\ufeff', '')
    content_text = content_text.replace('\r\n', '\n').replace('\r', '\n')
    content_text = content_text.replace('\u2028', '\n').replace('\u2029', '\n')

    # 如果整段文本被转义成了字面量 \n，恢复为真实换行。
    if '\n' not in content_text and '\\n' in content_text:
        content_text = content_text.replace('\\n', '\n')

    return content_text


def is_plain_document_title(line):
    """判断无 Markdown 标记时首个非空行是否像公文标题。"""
    stripped = line.strip()
    if not stripped or len(stripped) > 80:
        return False
    if stripped.endswith(('。', '；', ';', '：', ':')):
        return False
    title_suffixes = ('通知', '报告', '请示', '批复', '函', '意见', '方案', '总结', '纪要', '倡议书')
    return stripped.endswith(title_suffixes)


def promote_plain_title(lines):
    """没有 # 标题时，将首个疑似公文标题的非空行提升为标题。"""
    for idx, line in enumerate(lines):
        if not line.strip():
            continue
        if is_plain_document_title(line):
            lines[idx] = f"# {line.strip()}"
        return lines
    return lines


def is_legacy_ai_disclaimer(line):
    """过滤模型自行追加的旧版 AI 提示，统一由脚本在末尾生成规范提示。"""
    stripped = line.strip()
    return stripped in {
        "---",
        "内容由AI生成，仅供参考",
        "内容由 AI 生成，仅供参考",
        AI_DISCLAIMER_TEXT,
    }


def is_recipient(line):
    """判断是否为主送机关
    
    主送机关特征：
    - 包含：政府、厅、局、委、办公室、机构等
    - 以逗号、顿号分隔多个单位
    - 以冒号结尾
    """
    stripped = line.strip()
    
    # 必须以冒号结尾
    if not stripped.endswith('：') and not stripped.endswith(':'):
        return False

    # 排除明显不是主送机关的句子
    # 如："现就...通知如下："、"特此通知："等
    exclude_patterns = [
        r'通知如下：?$',
        r'决定如下：?$',
        r'批复如下：?$',
        r'意见如下：?$',
        r'复函如下：?$',
        r'函复如下：?$',
        r'报告如下：?$',
        r'请示如下：?$',
    ]
    if any(re.search(p, stripped) for p in exclude_patterns):
        return False

    # 占位示例或短主送机关也应识别，如“×××：”
    if len(stripped) <= 30:
        return True

    # 包含主送机关关键词
    keywords = ['政府', '厅', '局', '委', '办公室', '机构', '中心', '公司', '集团']
    has_keyword = any(kw in stripped for kw in keywords)
    
    if not has_keyword:
        return False
    
    # 排除过长的句子（主送机关通常不超过50字）
    if len(stripped) > 50:
        return False
    
    return True


def validate_input(content_text: str) -> bool:
    """验证输入内容"""
    if len(content_text) > MAX_INPUT_LENGTH:
        raise ValueError(f"输入内容过长（{len(content_text):,} 字符）")
    lines = content_text.split('\n')
    if len(lines) > MAX_INPUT_LINES:
        raise ValueError(f"输入行数过多（{len(lines):,} 行）")
    return True


def fix_reference_format(content_text: str) -> str:
    """修复引用格式"""
    def replace_ref(match):
        ref_num = match.group(1)
        return f'[^{ref_num}^]'
    content_text = re.sub(r'\[\^(\d+)\](?!\^)', replace_ref, content_text)
    content_text = re.sub(r'\[\^(\d+)\](?=\[\^\d+\])', replace_ref, content_text)
    return content_text


def parse_markdown_table(lines, start_idx):
    """
    从指定行开始解析Markdown表格，返回 (table_rows, next_idx)。
    
    支持格式：
    | 列1 | 列2 | 列3 |
    |-----|-----|-----|
    | 数据 | 数据 | 数据 |
    
    Returns:
        list[list[str]]: 解析后的单元格内容（不含分隔行）
        int: 表格结束后的下一行索引；若start_idx不是表格行则返回 ([], start_idx)
    """
    if start_idx >= len(lines):
        return [], start_idx
    
    stripped = lines[start_idx].strip()
    # 表格行必须以 | 开头（可选前导空格）
    if not stripped.startswith('|'):
        return [], start_idx
    
    def split_table_row(line):
        """拆分表格行，去掉首尾 | 并分割"""
        s = line.strip()
        if s.startswith('|'):
            s = s[1:]
        if s.endswith('|'):
            s = s[:-1]
        return [cell.strip() for cell in s.split('|')]
    
    def is_separator(line):
        """判断是否为分隔行（如 |---|---|---|）"""
        s = line.strip()
        if not s.startswith('|'):
            return False
        cells = split_table_row(s)
        return all(re.match(r'^[-:]+$', c) for c in cells)
    
    rows = []
    i = start_idx
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped.startswith('|'):
            break
        if is_separator(stripped):
            # 分隔行：记录但不加入数据行（已跳过）
            i += 1
            continue
        rows.append(split_table_row(stripped))
        i += 1
    
    # 至少要有2行（标题行 + 数据行）才算有效表格
    if len(rows) < 2:
        return [], start_idx
    
    return rows, i


def add_word_table(doc, table_rows, body_font, body_size):
    """
    将解析后的表格数据添加为Word原生表格。
    
    首行作为表头（加粗），其余为数据行。
    所有单元格使用公文字体。
    
    Args:
        doc: Document 对象
        table_rows: list[list[str]], 解析后的表格数据
        body_font: 正文字体名
        body_size: 正文字号(pt)
    """
    num_cols = max(len(row) for row in table_rows)
    
    table = doc.add_table(rows=len(table_rows), cols=num_cols)
    table.style = 'Table Grid'
    
    for row_idx, row_data in enumerate(table_rows):
        for col_idx in range(num_cols):
            cell = table.rows[row_idx].cells[col_idx]
            # 清空默认段落
            cell.paragraphs[0].clear()
            text = row_data[col_idx] if col_idx < len(row_data) else ''
            add_formatted_text(cell.paragraphs[0], text, body_font, body_size, bold=(row_idx == 0))
            set_paragraph_format(cell.paragraphs[0], first_line_indent=False, line_spacing=body_size + 12)
    
    # 表格后加一个空段落，与后续内容保持间距
    spacer = doc.add_paragraph()
    set_paragraph_format(spacer, first_line_indent=False, line_spacing=body_size + 12)


def set_paragraph_mark_font(para, font_name, font_size):
    """设置空段落的段落标记字体，避免空行高度受默认样式影响。"""
    pPr = para._p.get_or_add_pPr()
    rPr = pPr.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        pPr.append(rPr)

    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    latin_font = get_latin_font_config()
    rFonts.set(qn('w:ascii'), latin_font)
    rFonts.set(qn('w:hAnsi'), latin_font)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)

    for tag in ('w:sz', 'w:szCs'):
        elem = rPr.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            rPr.append(elem)
        elem.set(qn('w:val'), str(int(font_size * 2)))


def add_blank_paragraphs(doc, count, line_spacing, font_name=None, font_size=None):
    """插入指定数量的空段，用于公文固定空行。"""
    for _ in range(count):
        para = doc.add_paragraph()
        set_paragraph_format(para, first_line_indent=False, line_spacing=line_spacing)
        if font_name and font_size:
            set_paragraph_mark_font(para, font_name, font_size)


def create_document(content_text, output_path=None):
    """
    创建排版后的普通公文文档（无红头）
    
    Args:
        content_text: 正文内容（Markdown格式）
        output_path: 输出文件路径
    
    Returns:
        输出文件路径
    """
    content_text = normalize_content_text(content_text)
    validate_input(content_text)
    content_text = fix_reference_format(content_text)
    
    doc = Document()
    
    # 设置 A4 纸张和页边距（公文格式规范）
    page_width_mm, page_height_mm = get_page_size()
    margin_top, margin_bottom, margin_left, margin_right = get_page_margin()
    for section in doc.sections:
        section.page_width = Mm(page_width_mm)
        section.page_height = Mm(page_height_mm)
        section.top_margin = Cm(margin_top / 10)
        section.bottom_margin = Cm(margin_bottom / 10)
        section.left_margin = Cm(margin_left / 10)
        section.right_margin = Cm(margin_right / 10)
        add_center_page_number(section)
    
    # 处理正文内容
    lines = content_text.strip().split('\n')
    if not has_markdown_title(lines):
        inferred_title = infer_title_from_output_path(output_path)
        before = list(lines)
        lines = promote_plain_title(lines)
        if before == lines and inferred_title:
            lines = [f"# {inferred_title}", ""] + lines
    i = 0
    line_count = len(lines)
    
    # 从配置获取字体信息
    title_font, title_size, _ = get_font_config('title')
    title_line_spacing = FORMAT_CONFIG['title'].get('line_spacing_pt', 33) if FORMAT_CONFIG and 'title' in FORMAT_CONFIG else 33
    body_font, body_size, _ = get_font_config('body')
    h1_font, h1_size, _ = get_font_config('heading1')
    h2_font, h2_size, _ = get_font_config('heading2')
    h3_font, h3_size, h3_bold = get_font_config('heading3')
    body_line_spacing = FORMAT_CONFIG['body']['line_spacing_pt'] if FORMAT_CONFIG else 28
    
    # 附录区域标志：素材使用情况、参考来源和知识专库链接统一靠左排版
    in_appendix_section = False
    # 知识专库链接区域标志：URL 生成可点击蓝链
    in_kb_section = False
    # 附件正文首页“附件”后的下一行作为附件标题处理
    awaiting_attachment_title = False
    # 正文末尾附件目录后，落款前需要固定空三行
    attachment_list_pending_sign_gap = False

    while i < line_count:
        line = lines[i]
        stripped = line.strip()
        line_type = get_line_type(stripped)

        if is_legacy_ai_disclaimer(stripped):
            i += 1
            continue

        # Markdown表格检测（必须在其他判断之前）
        if stripped.startswith('|'):
            table_rows, next_i = parse_markdown_table(lines, i)
            if table_rows:
                add_word_table(doc, table_rows, body_font, body_size)
                i = next_i
                continue
            # 不是有效表格，当普通文本处理，继续往下走

        # 空行
        if line_type == 'empty':
            i += 1
            continue

        if (
            attachment_list_pending_sign_gap
            and not is_attachment_line(stripped)
            and not is_attachment_continuation(stripped)
        ):
            if is_signing_entity(stripped):
                add_blank_paragraphs(doc, 3, body_line_spacing, body_font, body_size)
            elif is_date_line(stripped):
                add_blank_paragraphs(doc, 3, body_line_spacing, body_font, body_size)
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                add_formatted_text(para, stripped, body_font, body_size)
                set_paragraph_format(para, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.RIGHT, line_spacing=body_line_spacing)
                attachment_list_pending_sign_gap = False
                i += 1
                continue
            attachment_list_pending_sign_gap = False

        if awaiting_attachment_title:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_formatted_text(para, stripped, title_font, title_size)
            set_paragraph_format(para, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=title_line_spacing)
            set_outline_level(para, 0)
            awaiting_attachment_title = False
            i += 1
            continue

        marker_text = strip_markdown_heading(stripped)
        if marker_text in ("素材使用情况", "参考资料", "参考来源"):
            in_appendix_section = True
            para = doc.add_paragraph()
            label = "【素材使用情况】" if marker_text == "素材使用情况" else f"【{marker_text}】"
            add_formatted_text(para, label, body_font, body_size)
            set_paragraph_format(para, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=body_line_spacing)
            i += 1
            continue

        if marker_text == "知识专库链接":
            in_appendix_section = True
            in_kb_section = True
            para = doc.add_paragraph()
            add_formatted_text(para, "【知识专库链接】", body_font, body_size)
            set_paragraph_format(para, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=body_line_spacing)
            i += 1
            continue

        # 检测附录区域开始
        if (
            re.match(r'^【素材使用情况】', stripped)
            or re.match(r'^【参考资料】', stripped)
            or re.match(r'^【参考来源】', stripped)
            or stripped in ("素材使用情况", "参考资料", "参考来源")
        ):
            in_appendix_section = True

        # 检测【知识专库链接】区域开始
        if re.match(r'^【知识专库链接】', stripped):
            in_appendix_section = True
            in_kb_section = True

        # 附件清单续行要先于三级标题识别，避免 “2.附件名” 被误判为三级标题。
        if i > 0 and is_attachment_continuation(stripped):
            prev_non_empty = ""
            for prev_idx in range(i - 1, -1, -1):
                prev_non_empty = lines[prev_idx].strip()
                if prev_non_empty:
                    break
            if is_attachment_line(prev_non_empty) or is_attachment_continuation(prev_non_empty):
                para = doc.add_paragraph()
                add_formatted_text(para, f"        {stripped}", body_font, body_size)
                set_attachment_list_format(para, continuation=True, line_spacing=body_line_spacing)
                attachment_list_pending_sign_gap = True
                i += 1
                continue

        # 公文标题（# 标题）
        if line_type == 'title':
            content = strip_markdown_heading(stripped)
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_formatted_text(para, content, title_font, title_size)
            set_paragraph_format(para, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=title_line_spacing)
            set_outline_level(para, 0)  # 文档标题设为1级
            add_blank_paragraphs(doc, 1, body_line_spacing, body_font, body_size)
            i += 1
            continue

        # 一级标题（## 一、XXX）
        if line_type == 'heading1':
            content = strip_markdown_heading(stripped)
            para = doc.add_paragraph()
            add_formatted_text(para, content, h1_font, h1_size)
            set_paragraph_format(para, indent_chars=2, line_spacing=body_line_spacing)
            set_outline_level(para, 1)  # 一级标题设为2级
            i += 1
            continue

        # 二级标题（### （一）XXX）
        if line_type == 'heading2':
            content = strip_markdown_heading(stripped)
            para = doc.add_paragraph()
            add_formatted_text(para, content, h2_font, h2_size)
            set_paragraph_format(para, indent_chars=2, line_spacing=body_line_spacing)
            set_outline_level(para, 2)  # 二级标题设为3级
            i += 1
            continue

        # 三级标题（#### 1.XXX）
        if line_type == 'heading3':
            content = strip_markdown_heading(stripped)
            para = doc.add_paragraph()
            add_formatted_text(para, content, h3_font, h3_size, bold=h3_bold)
            set_paragraph_format(para, indent_chars=2, line_spacing=body_line_spacing)
            set_outline_level(para, 3)  # 三级标题设为4级
            i += 1
            continue

        # 附件
        if is_attachment_line(stripped):
            add_blank_paragraphs(doc, 1, body_line_spacing, body_font, body_size)
            para = doc.add_paragraph()
            add_formatted_text(para, stripped, body_font, body_size)
            set_attachment_list_format(para, continuation=False, line_spacing=body_line_spacing)
            attachment_list_pending_sign_gap = True
            i += 1
            continue

        if is_attachment_section_title(stripped):
            if document_has_visible_content(doc):
                page_break_para = doc.add_paragraph()
                page_break_run = page_break_para.add_run()
                page_break_run.add_break(WD_BREAK.PAGE)
                set_paragraph_format(page_break_para, first_line_indent=False, line_spacing=body_line_spacing)
            para = doc.add_paragraph()
            add_formatted_text(para, stripped, h1_font, h1_size)
            set_paragraph_format(para, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=body_line_spacing)
            awaiting_attachment_title = True
            i += 1
            continue

        # 落款单位
        if is_signing_entity(stripped):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_formatted_text(para, stripped, body_font, body_size)
            set_paragraph_format(para, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.RIGHT, line_spacing=body_line_spacing)
            i += 1
            continue

        # 联系信息
        if is_contact_info(stripped):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_formatted_text(para, stripped, body_font, body_size)
            set_paragraph_format(para, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.RIGHT, line_spacing=body_line_spacing)
            i += 1
            continue

        # 落款日期：优先根据前一行落款单位和后续附录标记判断，避免素材页影响位置比例
        if should_right_align_date(lines, i):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_formatted_text(para, stripped, body_font, body_size)
            set_paragraph_format(para, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.RIGHT, line_spacing=body_line_spacing)
            i += 1
            continue

        # 分页符标记
        if stripped == '[分页符]':
            para = doc.add_paragraph()
            run = para.add_run()
            run.add_break(WD_BREAK.PAGE)
            i += 1
            continue

        # 主送机关（前10行内）
        if is_recipient(stripped) and i < 10:
            para = doc.add_paragraph()
            add_formatted_text(para, stripped, body_font, body_size)
            set_paragraph_format(para, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=body_line_spacing)
            i += 1
            continue

        # 附录区域：统一靠左排版；知识专库 URL 生成可点击蓝链
        if in_appendix_section:
            if in_kb_section and re.search(r'https?://', stripped):
                url_match = re.search(r'(https?://\S+)', stripped)
                if url_match:
                    url = sanitize_url(url_match.group(1))
                    if is_internal_search_endpoint(url):
                        i += 1
                        continue
                    para = doc.add_paragraph()
                    add_hyperlink(para, url, url, font_name=body_font, font_size=body_size)
                    set_paragraph_format(para, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=body_line_spacing)
                    i += 1
                    continue

            para = doc.add_paragraph()
            clean_text = strip_markdown_bold(stripped)
            add_formatted_text(para, clean_text, body_font, body_size)
            set_paragraph_format(para, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=body_line_spacing)
            i += 1
            continue

        # 知识专库链接区域：靠左排版，URL生成可点击蓝链
        if in_kb_section:
            # URL行：生成超链接
            if re.search(r'https?://', stripped):
                # 提取URL
                url_match = re.search(r'(https?://\S+)', stripped)
                if url_match:
                    url = sanitize_url(url_match.group(1))
                    if is_internal_search_endpoint(url):
                        i += 1
                        continue
                    para = doc.add_paragraph()
                    add_hyperlink(para, url, url, font_name=body_font, font_size=body_size)
                    set_paragraph_format(para, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=body_line_spacing)
                    i += 1
                    continue
            # 非URL行（标签、标题等）：靠左排版，不缩进
            para = doc.add_paragraph()
            clean_text = strip_markdown_bold(stripped)
            add_formatted_text(para, clean_text, body_font, body_size)
            set_paragraph_format(para, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=body_line_spacing)
            i += 1
            continue

        # 普通正文
        # 清理Markdown加粗标记
        clean_text = strip_markdown_bold(stripped)
        para = doc.add_paragraph()
        add_formatted_text(para, clean_text, body_font, body_size)
        set_paragraph_format(para, indent_chars=2, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=body_line_spacing)
        i += 1
    
    # 保存文件
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    if not output_path:
        output_path = f'{OUTPUT_DIR}/公文_{timestamp}.docx'

    if not output_path.lower().endswith('.docx'):
        output_path = output_path.rsplit('.', 1)[0] + '.docx' if '.' in output_path else output_path + '.docx'

    # 添加版本号（如果文件已存在）
    output_path = get_next_version(output_path)

    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)

    add_ai_disclaimer(doc)
    doc.save(output_path)
    return output_path


def main():
    """命令行入口"""
    parser = argparse.ArgumentParser(description='公文排版工具 v2.0.0 - 只支持普通格式')
    parser.add_argument('input', nargs='?', help='输入文件路径')
    parser.add_argument('--output', help='输出文件路径')
    parser.add_argument('--text', help='直接输入公文文本（如果是文件路径会自动读取）')
    
    args = parser.parse_args()
    
    if args.text:
        # 智能检测：如果 --text 是文件路径，自动读取文件内容
        if os.path.exists(args.text) and os.path.isfile(args.text):
            print(f'⚠ 检测到 --text 参数是文件路径，自动读取文件内容: {args.text}')
            with open(args.text, 'r', encoding='utf-8') as f:
                content = f.read()
        else:
            content = args.text
    elif args.input and os.path.exists(args.input):
        with open(args.input, 'r', encoding='utf-8') as f:
            content = f.read()
    else:
        parser.print_help()
        sys.exit(1)
    
    try:
        output_path = create_document(
            content,
            output_path=args.output
        )
        
        print(f'✓ 普通公文已生成: {output_path}')
    except Exception as e:
        print(f'✗ 生成失败: {e}')
        import traceback
        traceback.print_exc()


if __name__ == '__main__':
    main()
