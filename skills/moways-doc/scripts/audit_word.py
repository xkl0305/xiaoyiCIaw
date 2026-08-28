#!/usr/bin/env python3
"""
博维咨询 Word 报告排版审查脚本
依据《博维咨询Word版报告行文排版标准 V2.6》

用法:
    python audit_word.py input.docx [--output report.json] [--strict]

输出一份排版问题清单，按严重程度排序：
  - ERROR:   严重违规，必须修改
  - WARNING: 一般问题，建议修改
  - INFO:    提示信息，可酌情处理
"""

import argparse
import json
import os
import re
import sys

try:
    from docx import Document
    from docx.shared import Pt, Mm, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("错误: 需要安装 python-docx")
    print("运行: pip install python-docx --break-system-packages")
    sys.exit(1)


# 导入format_word中的常量和辅助函数
script_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, script_dir)
from format_word import (
    PAGE_TOP_MARGIN, PAGE_BOTTOM_MARGIN, PAGE_LEFT_MARGIN, PAGE_RIGHT_MARGIN,
    FONT_FANGSONG, FONT_KAITI, FONT_HEITI, FONT_SONGTI, FONT_TNR,
    BODY_FONT_SIZE, BODY_LINE_SPACING,
    H1_FONT, H1_BOLD, H2_FONT, H2_BOLD,
    TABLE_HEADER_FONT, TABLE_BODY_FONT,
    detect_heading_level, is_table_title, is_figure_title,
    RE_H1, RE_H2, RE_H3, RE_H4,
)
# v2.7 新增：中文标点国标 GB/T 15834—2011 检查
from punctuation import audit_chinese_punctuation


class Issue:
    """一条审查问题"""
    def __init__(self, severity, category, location, message, suggestion=''):
        self.severity = severity      # ERROR / WARNING / INFO
        self.category = category      # page_setup / font / heading / spacing / ...
        self.location = location      # 段落编号或描述
        self.message = message
        self.suggestion = suggestion

    def to_dict(self):
        return {
            'severity': self.severity,
            'category': self.category,
            'location': self.location,
            'message': self.message,
            'suggestion': self.suggestion,
        }

    def __str__(self):
        sev_icon = {'ERROR': '❌', 'WARNING': '⚠️', 'INFO': 'ℹ️'}.get(self.severity, '•')
        s = f"{sev_icon} [{self.severity}] {self.location}: {self.message}"
        if self.suggestion:
            s += f"\n   → 建议: {self.suggestion}"
        return s


def get_font_name(run):
    """获取run的中文字体名称"""
    rpr = run._element.find(qn('w:rPr'))
    if rpr is not None:
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is not None:
            ea = rfonts.get(qn('w:eastAsia'))
            if ea:
                return ea
    if run.font.name:
        return run.font.name
    return None


def get_font_size_pt(run):
    """获取run的字号（pt）"""
    if run.font.size:
        return run.font.size.pt
    return None


def get_line_spacing(para):
    """获取段落行距（倍数）"""
    pf = para.paragraph_format
    if pf.line_spacing is not None:
        try:
            return float(pf.line_spacing)
        except (TypeError, ValueError):
            pass
    return None


def has_explicit_page_break(para):
    """判断段落是否通过段前分页或手动分页符另起页。"""
    if para.paragraph_format.page_break_before:
        return True
    for run in para.runs:
        for child in run._element:
            if child.tag == qn('w:br') and child.get(qn('w:type')) == 'page':
                return True
    return False


def get_outline_level_value(para):
    """读取段落大纲级别；返回字符串值，如 '0'、'1'，无则返回 None。"""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        return None
    outline = pPr.find(qn('w:outlineLvl'))
    if outline is None:
        return None
    return outline.get(qn('w:val'))


def next_content_element_after_para(para):
    """返回段落后的下一个实质内容元素，跳过空段落。"""
    elem = para._element.getnext()
    while elem is not None:
        tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        if tag == 'p':
            texts = elem.findall('.//' + qn('w:t'))
            text = ''.join(t.text or '' for t in texts).strip()
            if text:
                return elem
        elif tag == 'tbl':
            return elem
        elem = elem.getnext()
    return None


def audit_document(input_path, strict=False):
    """
    审查文档排版，返回问题列表。

    Args:
        input_path: docx文件路径
        strict: 是否严格模式（更多检查项）

    Returns:
        list[Issue]: 问题列表，按严重程度排序
    """
    doc = Document(input_path)
    issues = []

    # ============================
    # 1. 页面设置检查
    # ============================
    for i, section in enumerate(doc.sections):
        loc = f"节{i+1}" if len(doc.sections) > 1 else "页面设置"

        # 页边距（允许1mm误差）
        tolerance = Mm(1)

        if section.top_margin is not None:
            diff = abs(section.top_margin - PAGE_TOP_MARGIN)
            if diff > tolerance:
                actual_mm = section.top_margin / Mm(1) if section.top_margin else 0
                issues.append(Issue(
                    'ERROR', 'page_setup', loc,
                    f'上边距 {actual_mm:.1f}mm，标准要求 37mm±1mm',
                    '设置上边距为37mm'
                ))

        if section.left_margin is not None:
            diff = abs(section.left_margin - PAGE_LEFT_MARGIN)
            if diff > tolerance:
                actual_mm = section.left_margin / Mm(1) if section.left_margin else 0
                issues.append(Issue(
                    'ERROR', 'page_setup', loc,
                    f'左边距（订口）{actual_mm:.1f}mm，标准要求 28mm±1mm',
                    '设置左边距为28mm'
                ))

    # ============================
    # 1b. 封面 / 目录 / 正文分页边界检查
    # ============================
    non_empty = [
        (i, p, p.text.strip())
        for i, p in enumerate(doc.paragraphs)
        if p.text.strip()
    ]
    toc_idx = None
    first_h1_idx = None
    for i, para, text in non_empty:
        if toc_idx is None and text in ('目录', '目 录'):
            toc_idx = i
        outline_level = get_outline_level_value(para)
        looks_like_toc_entry = '\t' in text
        if (first_h1_idx is None and
                (outline_level == '0' or
                 (RE_H1.match(text) and not looks_like_toc_entry))):
            first_h1_idx = i
        if toc_idx is not None and first_h1_idx is not None:
            break

    if toc_idx is not None:
        has_cover_like_content = any(i < toc_idx for i, _, _ in non_empty)
        if has_cover_like_content and not has_explicit_page_break(doc.paragraphs[toc_idx]):
            issues.append(Issue(
                'ERROR', 'pagination', f'第{toc_idx+1}段 "目录"',
                '目录未设置另起页，容易被排在封面底部',
                '给"目录"段落设置 page_break_before=True，或在封面末尾插入手动分页符'
            ))

        if first_h1_idx is not None and first_h1_idx > toc_idx:
            if not has_explicit_page_break(doc.paragraphs[first_h1_idx]):
                issues.append(Issue(
                    'ERROR', 'pagination', f'第{first_h1_idx+1}段',
                    '目录之后的正文一级标题未设置另起页',
                    '给正文第一个一级标题设置 page_break_before=True，确保正文从目录后一页开始'
                ))

    # 表题必须和后续表格绑定，避免表题独占上一页、表格另起下一页
    for i, para, text in non_empty:
        if is_table_title(text):
            next_elem = next_content_element_after_para(para)
            next_tag = None
            if next_elem is not None:
                next_tag = next_elem.tag.split('}')[-1] if '}' in next_elem.tag else next_elem.tag
            if next_tag != 'tbl':
                issues.append(Issue(
                    'ERROR', 'pagination', f'第{i+1}段 "{text[:30]}"',
                    '表标题后未紧跟表格，容易出现表题与表格分页分离',
                    '删除表题与表格之间的多余正文/分页/分节，或让排版脚本重新识别并绑定表题和表格'
                ))
            elif not para.paragraph_format.keep_with_next:
                issues.append(Issue(
                    'WARNING', 'pagination', f'第{i+1}段 "{text[:30]}"',
                    '表标题未设置与下段同页',
                    '给表标题设置 keep_with_next=True，确保表题不与表格跨页分离'
                ))

    # ============================
    # 2. 逐段落检查
    # ============================
    heading_sequence = {1: 0, 2: 0, 3: 0, 4: 0}
    prev_heading_level = 0
    table_count = 0
    figure_count = 0

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        loc = f"第{i+1}段"
        text_preview = text[:40] + ('...' if len(text) > 40 else '')
        loc_detail = f'{loc} "{text_preview}"'

        level = detect_heading_level(text)

        # ---- 标题检查 ----
        if level is not None:
            # 标题层级跳跃检查
            if level > prev_heading_level + 1 and prev_heading_level > 0:
                issues.append(Issue(
                    'WARNING', 'heading', loc_detail,
                    f'标题层级从{prev_heading_level}级跳到{level}级',
                    f'检查是否缺少{prev_heading_level+1}级标题'
                ))
            prev_heading_level = level

            # 标题字体检查
            for run in para.runs:
                font_name = get_font_name(run)
                if font_name and level == 1 and font_name not in [FONT_HEITI, '黑体']:
                    issues.append(Issue(
                        'ERROR', 'font', loc_detail,
                        f'一级标题字体为"{font_name}"，标准要求黑体',
                        f'将字体改为黑体'
                    ))
                elif font_name and level == 2 and font_name not in [FONT_KAITI, '楷体', '楷体_GB2312']:
                    issues.append(Issue(
                        'ERROR', 'font', loc_detail,
                        f'二级标题字体为"{font_name}"，标准要求楷体',
                        f'将字体改为楷体'
                    ))
                break  # 只检查第一个run

            # 标题末尾标点检查
            if level <= 2:
                # 一二级标题独占行，末尾不加标点
                if text[-1] in '。，；：！？,.;:!?':
                    issues.append(Issue(
                        'WARNING', 'punctuation', loc_detail,
                        f'{level}级标题末尾有标点符号"{text[-1]}"',
                        '独占行的标题末尾不应加标点'
                    ))

            # 标题不应首行缩进
            if para.paragraph_format.first_line_indent:
                indent_cm = para.paragraph_format.first_line_indent / Cm(1)
                if indent_cm > 0.1:
                    issues.append(Issue(
                        'WARNING', 'indent', loc_detail,
                        f'标题有首行缩进 {indent_cm:.2f}cm',
                        '标题不应设置首行缩进'
                    ))

        # ---- 正文检查 ----
        elif not is_table_title(text) and not is_figure_title(text):
            # 正文字体检查
            for run in para.runs:
                font_name = get_font_name(run)
                size_pt = get_font_size_pt(run)

                if font_name and font_name not in [
                    FONT_FANGSONG, '仿宋_GB2312', FONT_TNR, FONT_SONGTI,
                    FONT_KAITI, '楷体_GB2312', FONT_HEITI
                ]:
                    if strict:
                        issues.append(Issue(
                            'INFO', 'font', loc_detail,
                            f'正文使用了非标准字体"{font_name}"',
                            f'正文应使用仿宋体（中文）或Times New Roman（英文数字）'
                        ))
                break

            # 首行缩进检查（正文应有首行缩进2字符）
            indent = para.paragraph_format.first_line_indent
            if indent is not None and indent > 0:
                pass  # 有缩进，正常
            elif indent is None or indent == 0:
                # 检查是否用空格冒充缩进
                if text.startswith('  ') or text.startswith('　'):
                    issues.append(Issue(
                        'WARNING', 'indent', loc_detail,
                        '使用空格实现首行缩进',
                        '应使用段落格式的首行缩进（2个字符），禁止用空格做占位符'
                    ))

            # 行距检查
            line_sp = get_line_spacing(para)
            if line_sp is not None and abs(line_sp - 1.5) > 0.1:
                if strict:
                    issues.append(Issue(
                        'INFO', 'spacing', loc_detail,
                        f'行距为{line_sp}倍，标准要求1.5倍',
                        '设置行距为1.5倍'
                    ))

        # ---- 表格/图标题检查 ----
        if is_table_title(text):
            table_count += 1
            expected_num = table_count
            m = re.match(r'表\s*(\d+)', text)
            if m:
                actual_num = int(m.group(1))
                if actual_num != expected_num:
                    issues.append(Issue(
                        'WARNING', 'numbering', loc_detail,
                        f'表格序号为{actual_num}，期望为{expected_num}',
                        f'按全文出现顺序重新编号'
                    ))

        if is_figure_title(text):
            figure_count += 1
            expected_num = figure_count
            m = re.match(r'图\s*(\d+)', text)
            if m:
                actual_num = int(m.group(1))
                if actual_num != expected_num:
                    issues.append(Issue(
                        'WARNING', 'numbering', loc_detail,
                        f'图例序号为{actual_num}，期望为{expected_num}',
                        f'按全文出现顺序重新编号'
                    ))

    # ============================
    # 3. 标题编号规范检查
    # ============================
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # 检查三级标题是否用顿号而非下角点
        m = re.match(r'^(\d+)、', text)
        if m:
            issues.append(Issue(
                'ERROR', 'numbering', f'第{i+1}段 "{text[:30]}..."',
                f'三级标题序号"{m.group(0)}"使用了顿号',
                '阿拉伯数字序号后应用下角点，如"1."而非"1、"'
            ))

        # 检查理工科编号法（如 1.2.1）
        if re.match(r'^\d+\.\d+', text):
            issues.append(Issue(
                'INFO', 'numbering', f'第{i+1}段 "{text[:30]}..."',
                '使用了理工科类论文标题编码法（如"1.2.1"）',
                '提交甲方的报告一般不使用此编码法，建议使用"一、""（一）""1."层级'
            ))

    # ============================
    # 4. 标点符号检查（v2.7：依据 GB/T 15834—2011）
    # ============================
    # 将段落文本喂给共享的 punctuation 模块，覆盖：
    # P1-P6 中文上下文半角标点（逗号/句号/分号/冒号/问号/叹号）
    # P7-P8 ASCII 直引号包中文
    # P9    半角括号包中文
    # P10   尖括号冒充书名号
    # P11   双连字符冒充破折号
    # P12   单 EM dash 在中文上下文（INFO）
    # P13-P14 三个英文点/单 … 冒充省略号
    # P15-P16 数值范围半角连字符/波浪
    # P17-P18 书名号/引号间多余顿号
    # P19   三级标题序号顿号误用
    # P20   百分号前空格
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        # 跳过纯非中文段落（中文占比 < 10%）
        cjk_chars = len(re.findall(r'[一-鿿]', text))
        if cjk_chars / max(len(text), 1) < 0.1:
            continue
        loc = f'第{i+1}段'
        text_preview = text[:40] + ('...' if len(text) > 40 else '')
        loc_detail = f'{loc} "{text_preview}"'
        for pi in audit_chinese_punctuation(text):
            issues.append(Issue(
                pi.severity, f'punctuation/{pi.code}', loc_detail,
                pi.message + (f' [样本: {pi.sample}]' if pi.sample else ''),
                pi.suggestion
            ))
    # 表格单元格也扫一遍
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if not cell_text:
                    continue
                cjk_chars = len(re.findall(r'[一-鿿]', cell_text))
                if cjk_chars / max(len(cell_text), 1) < 0.1:
                    continue
                loc = f'表{ti+1}[{ri+1},{ci+1}]'
                for pi in audit_chinese_punctuation(cell_text):
                    issues.append(Issue(
                        pi.severity, f'punctuation/{pi.code}', loc,
                        pi.message + (f' [样本: {pi.sample}]' if pi.sample else ''),
                        pi.suggestion
                    ))

    # ============================
    # 5. 数字格式检查
    # ============================
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # 年份缩写检查
        if re.search(r"(?<!\d)\d{2}年", text) and not re.search(r'\d{4}年', text):
            issues.append(Issue(
                'WARNING', 'number_format', f'第{i+1}段',
                '年份可能使用了缩写（如"21年"）',
                '年份应标全称（如"2021年"）'
            ))

        # 百分比范围写法检查
        m = re.search(r'(\d+)[\~～—\-](\d+)%', text)
        if m:
            issues.append(Issue(
                'WARNING', 'number_format', f'第{i+1}段',
                f'百分比范围"{m.group(0)}"前一个数缺少%',
                '每个百分数的%都不能省略，如"15%～30%"'
            ))

    # ============================
    # 6. 机构名称检查
    # ============================
    # 检查是否有"以下简称"但后文未使用
    abbreviations = {}
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        m = re.search(r'以下简称["""](.+?)["""]', text)
        if m:
            abbreviations[m.group(1)] = i

    # ============================
    # 排序并返回
    # ============================
    severity_order = {'ERROR': 0, 'WARNING': 1, 'INFO': 2}
    issues.sort(key=lambda x: severity_order.get(x.severity, 3))

    return issues


def main():
    parser = argparse.ArgumentParser(
        description='博维咨询Word报告排版审查工具（依据V2.6标准）'
    )
    parser.add_argument('input', help='输入docx文件路径')
    parser.add_argument('--output', '-o', default=None,
                        help='输出审查报告JSON文件路径')
    parser.add_argument('--strict', action='store_true',
                        help='严格模式（更多检查项）')
    parser.add_argument('--only-punctuation', action='store_true',
                        help='只显示标点类问题（P1-P20）')

    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"错误: 文件不存在 - {args.input}")
        sys.exit(1)

    print(f"正在审查: {args.input}")
    issues = audit_document(args.input, strict=args.strict)
    if args.only_punctuation:
        issues = [i for i in issues if i.category.startswith('punctuation')]

    # 统计
    errors = sum(1 for i in issues if i.severity == 'ERROR')
    warnings = sum(1 for i in issues if i.severity == 'WARNING')
    infos = sum(1 for i in issues if i.severity == 'INFO')

    print(f"\n审查完成! 共发现 {len(issues)} 个问题:")
    print(f"  ❌ 错误: {errors}")
    print(f"  ⚠️  警告: {warnings}")
    print(f"  ℹ️  提示: {infos}")

    if issues:
        print(f"\n{'='*60}")
        for issue in issues:
            print(f"\n{issue}")

    # 输出JSON
    if args.output:
        report = {
            'file': args.input,
            'summary': {
                'total': len(issues),
                'errors': errors,
                'warnings': warnings,
                'infos': infos,
            },
            'issues': [i.to_dict() for i in issues],
        }
        with open(args.output, 'w', encoding='utf-8') as f:
            json.dump(report, f, ensure_ascii=False, indent=2)
        print(f"\n审查报告已保存: {args.output}")

    return len(issues)


if __name__ == '__main__':
    sys.exit(0 if main() == 0 else 1)
