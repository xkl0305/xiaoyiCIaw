#!/usr/bin/env python3
"""
博维咨询 Word 报告自动排版脚本
依据《博维咨询Word版报告行文排版标准 V2.6》

用法:
    python format_word.py input.docx [output.docx] [--mode full|light] [--cover] [--toc]

参数:
    input.docx   : 输入文件路径
    output.docx  : 输出文件路径（默认在原文件名后加 _formatted）
    --mode full   : 完整排版（默认），覆盖所有格式
    --mode light  : 轻量排版，仅修正字体/字号/行距，保留其他格式
    --cover       : 标记第一页为封面页（不加页码，不计入页码数）
    --toc         : 标记文档包含目录页（目录格式单独处理）
    --page-num-style dash : 页码样式 "— X —"（默认）
    --page-num-style plain : 页码样式 "X"

v2.7.3 变更：
    一级标题（章节）另起页自动处理——除第一个一级标题外，
    后续一级标题自动设置 page_break_before = True 强制另起页；
    二级及以下标题确保不强制分页。
"""

import argparse
import copy
import os
import re
import sys

try:
    from docx import Document
    from docx.shared import Pt, Mm, Cm, Emu, Inches
    from docx.shared import RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    import lxml.etree as etree
except ImportError:
    print("错误: 需要安装 python-docx 和 lxml")
    print("运行: pip install python-docx lxml --break-system-packages")
    sys.exit(1)

# v2.7 新增：中文标点国标 GB/T 15834—2011 自动修正
_script_dir = os.path.dirname(os.path.abspath(__file__))
if _script_dir not in sys.path:
    sys.path.insert(0, _script_dir)
from punctuation import fix_chinese_punctuation  # noqa: E402


# =============================================================================
# 博维排版标准常量（源自V2.6文档）
# =============================================================================

# 页面设置
PAGE_TOP_MARGIN = Mm(37)       # 天头 37mm
PAGE_BOTTOM_MARGIN = Mm(35)    # 下边距（版心225mm推算）
PAGE_LEFT_MARGIN = Mm(28)      # 订口 28mm
PAGE_RIGHT_MARGIN = Mm(26)     # 右边距（版心156mm推算: 210-28-156=26）

# 字体定义
FONT_FANGSONG = '仿宋'         # 正文中文字体
FONT_KAITI = '楷体'            # 二级标题 / 封面说明
FONT_HEITI = '黑体'            # 一级标题 / 表格标题
FONT_XIAOBAOSONG = '方正小标宋简体'  # 全文大标题（备选黑体）
FONT_SONGTI = '宋体'           # 页码字体
FONT_TNR = 'Times New Roman'   # 英文/数字字体

# 正文格式（适用于内容较多、需要目录的报告）
BODY_FONT_SIZE = Pt(12)        # 小四号 = 12pt
BODY_LINE_SPACING = 1.5        # 1.5倍行距
BODY_FIRST_LINE_INDENT = Cm(0.85)  # 首行缩进2个字符（小四号下约0.85cm）
BODY_PARA_BEFORE = Pt(0)       # 段前间距
BODY_PARA_AFTER = Pt(0)        # 段后间距

# 封面格式
COVER_TITLE_FONT = FONT_HEITI          # 用户要求黑体
COVER_TITLE_SIZE = Pt(22)              # 二号字 = 22pt
COVER_TITLE_BOLD = True
COVER_SUBTITLE_FONT = FONT_KAITI
COVER_SUBTITLE_SIZE = Pt(14)           # 四号字 = 14pt
COVER_AUTHOR_FONT = FONT_KAITI
COVER_AUTHOR_SIZE = Pt(14)             # 四号字

# 标题格式
# 一级标题：黑体，加粗，与正文字号一致
H1_FONT = FONT_HEITI
H1_BOLD = True
H1_SIZE = BODY_FONT_SIZE

# 二级标题：楷体，加粗，与正文字号一致
H2_FONT = FONT_KAITI
H2_BOLD = True
H2_SIZE = BODY_FONT_SIZE

# 三四级标题：与正文一致
H3_FONT = FONT_FANGSONG
H3_BOLD = False
H3_SIZE = BODY_FONT_SIZE

# 目录格式
TOC_TITLE_FONT = FONT_HEITI
TOC_TITLE_SIZE = Pt(22)        # 二号字
TOC_TITLE_BOLD = True
TOC_L1_FONT = FONT_HEITI
TOC_L1_SIZE = Pt(14)           # 四号字
TOC_L1_BOLD = True
TOC_L2_FONT = FONT_KAITI
TOC_L2_SIZE = Pt(12)           # 小四号
TOC_L2_BOLD = True
TOC_L3_FONT = FONT_HEITI
TOC_L3_SIZE = Pt(10.5)         # 五号字

# 页眉格式
HEADER_LEFT_FONT = FONT_KAITI
HEADER_LEFT_SIZE = Pt(10.5)    # 五号字
HEADER_LEFT_BOLD = True
HEADER_RIGHT_FONT = FONT_HEITI
HEADER_RIGHT_SIZE = Pt(9)      # 小五号

# 页码格式
PAGE_NUM_FONT = FONT_SONGTI
PAGE_NUM_SIZE = Pt(10.5)       # 五号字

# 表格格式
TABLE_TITLE_FONT = FONT_HEITI
TABLE_TITLE_SIZE = Pt(12)      # 小四号
TABLE_TITLE_BOLD = True
TABLE_HEADER_FONT = FONT_KAITI
TABLE_HEADER_SIZE = Pt(9)      # 小五号
TABLE_HEADER_BOLD = True
TABLE_BODY_FONT = FONT_SONGTI
TABLE_BODY_SIZE = Pt(9)        # 小五号

# 图例格式
FIGURE_TITLE_FONT = FONT_KAITI
FIGURE_TITLE_SIZE = Pt(12)     # 小四号


# =============================================================================
# 标题层次识别正则
# =============================================================================

# 一级标题: "一、" "二、" ... "十、" "十一、" ...
RE_H1 = re.compile(r'^[一二三四五六七八九十百]+、')
# 二级标题: "（一）" "（二）" ...
RE_H2 = re.compile(r'^（[一二三四五六七八九十百]+）')
# 三级标题: "1." "2." ... (阿拉伯数字+下角点)
RE_H3 = re.compile(r'^(\d+)\.')
# 四级标题: "（1）" "（2）" ...
RE_H4 = re.compile(r'^（(\d+)）')
# 五级标题: "①" "②" ...
RE_H5 = re.compile(r'^[①②③④⑤⑥⑦⑧⑨⑩]')

# 附件标题识别（v2.7.3）: "附件1" "附件2" "附件1：" "附件2：XXX"
# 注意：要求附件序号后紧跟冒号或直接结束（不匹配"附件1的内容"这种正文）
RE_APPENDIX = re.compile(r'^附件\s*\d+\s*[：:．.]|^附件\s*\d+\s*$')

# 图表标题识别
# 表标题常见写法包括："表1 XXX"、"表1-XXX"、"表1：XXX"、"表 1—XXX"。
# 宽表格横向分节依赖该识别，漏识别会导致表题和表格分页分离。
RE_TABLE_TITLE = re.compile(r'^表\s*\d+\s*(?:[ \t　：:．.\-—–]|$)')
RE_FIGURE_TITLE = re.compile(r'^图\s*\d+\s')


def detect_heading_level(text, para=None):
    """检测段落的标题层级，返回 1-5 或 None（正文）

    优先级：
    1. 通过文本正则匹配标准编号（一、（一）、1.、（1）、①）
    2. 若正则未匹配，回退到 Word 样式名（Heading 1~5）检测
       适用于非标准编号的文档（如法律协议用纯文字标题）
    """
    text = text.strip()
    if not text:
        return None
    if RE_H1.match(text):
        return 1
    if RE_H2.match(text):
        return 2
    if RE_H3.match(text):
        return 3
    if RE_H4.match(text):
        return 4
    if RE_H5.match(text):
        return 5
    # v2.7.3: 回退到 Word 样式名检测
    if para is not None:
        style_name = para.style.name if para.style else ''
        # 匹配 "Heading 1", "Heading 2", ... 以及中文 "标题 1", "标题 2"
        import re as _re
        m = _re.match(r'^(?:Heading|标题)\s*(\d+)$', style_name.strip())
        if m:
            lv = int(m.group(1))
            if 1 <= lv <= 5:
                return lv
    return None


def is_table_title(text):
    """检测是否为表格标题（如 '表3 拉萨市农贸市场一览表'）"""
    return bool(RE_TABLE_TITLE.match(text.strip()))


def is_figure_title(text):
    """检测是否为图例标题（如 '图3 拉萨市农贸市场分布图'）"""
    return bool(RE_FIGURE_TITLE.match(text.strip()))


def has_chinese(text):
    """检测文本是否包含中文字符"""
    return bool(re.search(r'[\u4e00-\u9fff]', text))


def is_mostly_number_or_english(text):
    """检测文本是否主要由数字/英文构成"""
    if not text.strip():
        return False
    chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
    total_chars = len(text.strip())
    return chinese_chars / max(total_chars, 1) < 0.3


# =============================================================================
# 页面设置
# =============================================================================

def setup_page(section):
    """设置页面尺寸和页边距（A4纵向）"""
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = PAGE_TOP_MARGIN
    section.bottom_margin = PAGE_BOTTOM_MARGIN
    section.left_margin = PAGE_LEFT_MARGIN
    section.right_margin = PAGE_RIGHT_MARGIN


# =============================================================================
# OOXML 元素顺序辅助函数
# =============================================================================

# pPr子元素的OOXML schema顺序（rPr和sectPr必须在最后）
_PPR_ORDER = [
    'pStyle', 'keepNext', 'keepLines', 'pageBreakBefore', 'framePr',
    'widowControl', 'numPr', 'suppressLineNumbers', 'pBdr', 'shd',
    'tabs', 'suppressAutoHyphens', 'kinsoku', 'wordWrap', 'overflowPunct',
    'topLinePunct', 'autoSpaceDE', 'autoSpaceDN', 'bidi', 'adjustRightInd',
    'snapToGrid', 'spacing', 'ind', 'contextualSpacing', 'mirrorIndents',
    'suppressOverlap', 'jc', 'textDirection', 'textAlignment',
    'textboxTightWrap', 'outlineLvl', 'divId', 'cnfStyle', 'rPr',
    'sectPr', 'pPrChange',
]

# style子元素的OOXML schema顺序
_STYLE_ORDER = [
    'name', 'aliases', 'basedOn', 'next', 'link', 'autoRedefine',
    'hidden', 'uiPriority', 'semiHidden', 'unhideWhenUsed', 'qFormat',
    'locked', 'personal', 'personalCompose', 'personalReply', 'rsid',
    'pPr', 'rPr', 'tblPr', 'trPr', 'tcPr', 'tblStylePr',
]


def _insert_in_order(parent, child, order_list):
    """按schema顺序将子元素插入到parent中的正确位置。

    如果child的tag已存在于parent中，则替换；否则找到正确位置插入。
    """
    child_tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

    # 找到child在顺序表中的位置
    try:
        child_pos = order_list.index(child_tag)
    except ValueError:
        # 不在顺序表中，追加到rPr之前（安全位置）
        parent.append(child)
        return

    # 找到第一个顺序大于child的已有元素，插入其前面
    for i, existing in enumerate(parent):
        existing_tag = existing.tag.split('}')[-1] if '}' in existing.tag else existing.tag
        try:
            existing_pos = order_list.index(existing_tag)
        except ValueError:
            continue
        if existing_pos > child_pos:
            parent.insert(i, child)
            return

    # 没有比child顺序更大的元素，追加到末尾
    parent.append(child)


# =============================================================================
# 字体设置辅助函数
# =============================================================================

def _clear_theme_fonts(rfonts):
    """清除rFonts元素上的所有主题字体属性。

    OOXML中主题字体属性(asciiTheme/eastAsiaTheme/hAnsiTheme/cstheme)
    优先级高于显式字体名(ascii/eastAsia/hAnsi/cs)。
    如果不清除，即使设置了eastAsia='仿宋'，Word仍会使用主题字体（通常是宋体）。
    """
    for attr in ('w:asciiTheme', 'w:eastAsiaTheme', 'w:hAnsiTheme', 'w:cstheme'):
        key = qn(attr)
        if key in rfonts.attrib:
            del rfonts.attrib[key]


def _set_run_element_font(r_elem, cn_font, en_font=FONT_TNR, size=None, bold=None, color=None):
    """直接操作w:r XML元素设置字体（不依赖python-docx的Run对象）。

    这样可以处理嵌套在hyperlink/sdt/ins/del等容器内的run，
    而python-docx的para.runs只能看到段落的直接子run。
    """
    rpr = r_elem.find(qn('w:rPr'))
    if rpr is None:
        rpr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
        r_elem.insert(0, rpr)

    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rpr.insert(0, rfonts)

    _clear_theme_fonts(rfonts)
    rfonts.set(qn('w:eastAsia'), cn_font)
    rfonts.set(qn('w:ascii'), en_font)
    # hAnsi 控制非ASCII非东亚字符（含中文引号""、书名号等标点），
    # 必须用中文字体，否则标点会用TNR渲染导致显示异常
    rfonts.set(qn('w:hAnsi'), cn_font)
    rfonts.set(qn('w:cs'), en_font)

    if size is not None:
        sz = rpr.find(qn('w:sz'))
        if sz is None:
            sz = parse_xml(f'<w:sz {nsdecls("w")}/>')
            rpr.append(sz)
        sz.set(qn('w:val'), str(int(size.pt * 2)))  # half-points
        sz_cs = rpr.find(qn('w:szCs'))
        if sz_cs is None:
            sz_cs = parse_xml(f'<w:szCs {nsdecls("w")}/>')
            rpr.append(sz_cs)
        sz_cs.set(qn('w:val'), str(int(size.pt * 2)))

    if bold is not None:
        b = rpr.find(qn('w:b'))
        if bold:
            if b is None:
                rpr.append(parse_xml(f'<w:b {nsdecls("w")}/>'))
        else:
            if b is not None:
                rpr.remove(b)
        b_cs = rpr.find(qn('w:bCs'))
        if bold:
            if b_cs is None:
                rpr.append(parse_xml(f'<w:bCs {nsdecls("w")}/>'))
        else:
            if b_cs is not None:
                rpr.remove(b_cs)

    if color is not None:
        c = rpr.find(qn('w:color'))
        if c is None:
            c = parse_xml(f'<w:color {nsdecls("w")}/>')
            rpr.append(c)
        c.set(qn('w:val'), str(color))


def set_run_font(run, cn_font, en_font=FONT_TNR, size=None, bold=None, color=None):
    """
    设置run的字体（python-docx Run对象版本）。
    内部委托给 _set_run_element_font。
    """
    _set_run_element_font(run._element, cn_font, en_font, size, bold, color)


def set_paragraph_font(para, cn_font, en_font=FONT_TNR, size=None, bold=None, color=None):
    """设置段落中所有run的字体（包括嵌套在hyperlink/sdt/ins/del等内的run）。
    
    v2.7.3 新增 color 参数：强制设置字体颜色（如 RGBColor(0,0,0) 黑色）。
    """
    p_elem = para._element

    # 用xpath找到段落内所有run（包括嵌套的），而非只用para.runs
    all_runs = p_elem.findall('.//' + qn('w:r'))

    if not all_runs and para.text.strip():
        # 段落有文本但无run：创建run
        run = para.add_run(para.text)
        for child in list(p_elem):
            if child.tag == qn('w:t'):
                p_elem.remove(child)
        all_runs = [run._element]

    for r_elem in all_runs:
        _set_run_element_font(r_elem, cn_font, en_font, size, bold, color)

    # 同时设置段落级别的rPr（影响段落标记的默认字体）
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is not None:
        rpr = pPr.find(qn('w:rPr'))
        if rpr is None:
            rpr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
            pPr.append(rpr)
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
            rpr.insert(0, rfonts)
        _clear_theme_fonts(rfonts)
        rfonts.set(qn('w:eastAsia'), cn_font)
        rfonts.set(qn('w:ascii'), en_font)
        rfonts.set(qn('w:hAnsi'), cn_font)
        # v2.7.3: 段落级 rPr 也设置颜色
        if color is not None:
            c = rpr.find(qn('w:color'))
            if c is None:
                c = parse_xml(f'<w:color {nsdecls("w")}/>')
                rpr.append(c)
            c.set(qn('w:val'), str(color))
    else:
        # 如果段落连pPr都没有，创建完整的pPr/rPr/rFonts
        color_attr = f' w:val="{color}"' if color is not None else ''
        color_elem = f'<w:color {nsdecls("w")}{color_attr}/>' if color is not None else ''
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}><w:rPr><w:rFonts w:eastAsia="{cn_font}" w:ascii="{en_font}" w:hAnsi="{cn_font}"/>{color_elem}</w:rPr></w:pPr>')
        p_elem.insert(0, pPr)


def _reset_document_default_fonts(doc):
    """重置文档级别的默认字体设置和排版规则。

    清除Normal样式和文档默认rPr中的主题字体属性，
    并将默认字体设为正文标准（仿宋 + TNR）。
    同时在Normal样式上启用东亚避头尾规则。
    """
    # 1. 修复Normal样式
    try:
        normal = doc.styles['Normal']
        # 在Normal样式的pPr上启用避头尾，让所有段落默认继承
        pPr = normal.element.find(qn('w:pPr'))
        if pPr is None:
            pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
            # 按schema顺序插入：pPr在name/qFormat/rsid之后
            _insert_in_order(normal.element, pPr, _STYLE_ORDER)
        for tag, val in [('w:kinsoku', '1'), ('w:overflowPunct', '1'),
                         ('w:autoSpaceDE', '1'), ('w:autoSpaceDN', '1')]:
            elem = pPr.find(qn(tag))
            if elem is None:
                elem = parse_xml(f'<{tag} {nsdecls("w")} w:val="{val}"/>')
                _insert_in_order(pPr, elem, _PPR_ORDER)
    except KeyError:
        pass

    try:
        normal = doc.styles['Normal']
        rpr = normal.element.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
            rpr.insert(0, rfonts)
        _clear_theme_fonts(rfonts)
        rfonts.set(qn('w:eastAsia'), FONT_FANGSONG)
        rfonts.set(qn('w:ascii'), FONT_TNR)
        rfonts.set(qn('w:hAnsi'), FONT_FANGSONG)  # 中文标点走hAnsi通道，必须用中文字体
        rfonts.set(qn('w:cs'), FONT_TNR)
        normal.font.size = BODY_FONT_SIZE
    except KeyError:
        pass

    # 2. 修复文档级别的docDefaults中的主题字体
    styles_element = doc.styles.element
    doc_defaults = styles_element.find(qn('w:docDefaults'))
    if doc_defaults is not None:
        rpr_default = doc_defaults.find(qn('w:rPrDefault'))
        if rpr_default is not None:
            rpr = rpr_default.find(qn('w:rPr'))
            if rpr is not None:
                rfonts = rpr.find(qn('w:rFonts'))
                if rfonts is not None:
                    _clear_theme_fonts(rfonts)
                    rfonts.set(qn('w:eastAsia'), FONT_FANGSONG)
                    rfonts.set(qn('w:ascii'), FONT_TNR)
                    rfonts.set(qn('w:hAnsi'), FONT_FANGSONG)

    # 3. 修复其他常见内置样式的主题字体
    for style_name in ('Body Text', 'List Paragraph', 'No Spacing',
                       'Default Paragraph Font'):
        try:
            s = doc.styles[style_name]
            rpr = s.element.find(qn('w:rPr'))
            if rpr is not None:
                rfonts = rpr.find(qn('w:rFonts'))
                if rfonts is not None:
                    _clear_theme_fonts(rfonts)
        except KeyError:
            pass


# 孤立序号段落的正则：仅包含序号本身（可能带少量空白），没有实质正文内容
RE_ORPHAN_H1 = re.compile(r'^[一二三四五六七八九十百]+、\s*$')
RE_ORPHAN_H2 = re.compile(r'^（[一二三四五六七八九十百]+）\s*$')
RE_ORPHAN_H3 = re.compile(r'^\d+\.\s*$')
RE_ORPHAN_H4 = re.compile(r'^（\d+）\s*$')
RE_ORPHAN_H5 = re.compile(r'^[①②③④⑤⑥⑦⑧⑨⑩]\s*$')
# 孤立的圆点/句点（残留的列表标记）
RE_ORPHAN_DOT = re.compile(r'^[.．·•]\s*$')

# 段落断裂模式：上一段以数字+连字符结尾（如 "3-"  "15-"），下一段以数字开头（如 "4个" "20人"）
RE_BROKEN_TAIL = re.compile(r'[\d]-\s*$')
RE_BROKEN_HEAD = re.compile(r'^\d')


def _is_orphan_heading(text):
    """判断段落是否为孤立序号或残留标记（序号独占一段，正文在下一段）"""
    t = text.strip()
    if not t:
        return False
    return bool(
        RE_ORPHAN_H1.match(t) or RE_ORPHAN_H2.match(t) or
        RE_ORPHAN_H3.match(t) or RE_ORPHAN_H4.match(t) or
        RE_ORPHAN_H5.match(t) or RE_ORPHAN_DOT.match(t)
    )


def _should_merge_broken(text_current, text_next):
    """判断两个相邻段落是否因断裂需要合并。

    检测模式：
    - 上段以"数字-"结尾（如"分为3-"），下段以数字开头（如"4个讨论小组"）
    - 上段以"；"以外的标点或无标点结尾，下段以"；"开头（续接分号列表）
    """
    if not text_current or not text_next:
        return False
    # 数字范围断裂：如 "3-\n4个" "15-\n20人"
    if RE_BROKEN_TAIL.search(text_current) and RE_BROKEN_HEAD.match(text_next):
        return True
    return False


def _is_junk_paragraph(text):
    """判断段落是否为无意义的残留标记（应直接删除而非合并）"""
    t = text.strip()
    return bool(RE_ORPHAN_DOT.match(t)) if t else False


def _merge_paragraphs(doc):
    """合并孤立序号段落和断裂段落，删除无意义残留段落。

    处理三类问题：
    1. 无意义残留：孤立圆点 "." 等 → 直接删除
    2. 孤立序号：段落仅含序号（如 "2." "（一）"），下一段是正文 → 合并
    3. 段落断裂：文本在数字范围中间断开（如 "3-" 换行 "4个"）→ 合并

    使用多轮扫描确保级联合并正确处理。
    """
    total_ops = 0

    for _pass in range(5):  # 最多5轮，防止无限循环
        body = doc.element.body
        # 每轮重新获取段落列表（因为上一轮可能修改了结构）
        paras = list(doc.paragraphs)
        ops_this_round = 0
        i = 0

        while i < len(paras):
            text = paras[i].text
            elem = paras[i]._element

            # 第1类：删除无意义残留段落（孤立圆点等）
            if _is_junk_paragraph(text):
                if elem.getparent() is not None:
                    elem.getparent().remove(elem)
                    ops_this_round += 1
                    paras.pop(i)
                    continue  # 不递增i

            # 需要有下一段才能做合并判断
            if i >= len(paras) - 1:
                break

            next_text = paras[i + 1].text.strip()
            should_merge = False

            # 第2类：孤立序号合并
            if _is_orphan_heading(text) and next_text:
                should_merge = True
            # 第3类：段落断裂合并
            elif _should_merge_broken(text.rstrip(), next_text):
                should_merge = True

            if should_merge:
                current_elem = elem
                next_elem = paras[i + 1]._element

                for child in list(next_elem):
                    if child.tag == qn('w:pPr'):
                        continue
                    current_elem.append(child)

                if next_elem.getparent() is not None:
                    next_elem.getparent().remove(next_elem)
                paras.pop(i + 1)
                ops_this_round += 1
                # 不递增i，因为合并后可能还需要继续合并
            else:
                i += 1

        total_ops += ops_this_round
        if ops_this_round == 0:
            break  # 本轮无操作，提前退出

    if total_ops > 0:
        print(f"  预处理: 处理了 {total_ops} 个断裂/孤立/残留段落")

    return total_ops


# 项目符号字符集（段首出现时应清除）
# 包括：半角点、全角点、中圆点、实心圆、空心圆、方块、菱形、短横线等
RE_BULLET_PREFIX = re.compile(r'^[.．·•●○■□◆◇\-–—]\s*')


def _clean_bullet_markers(doc):
    """清除段落中的项目符号标记。

    处理两类问题：
    1. 文本字符型项目符号：段首的 "." "·" "•" "■" 等字符 → 从文本中移除
    2. Word列表编号(numPr)：让段落脱离编号列表 → 删除numPr属性

    注意：不清除博维标准的合法序号（如"一、" "（一）" "1." "（1）" "①"），
    这些由标题识别逻辑处理。
    """
    cleaned = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 跳过合法标题序号（不要误清除"1."这种三级标题）
        if detect_heading_level(text) is not None:
            continue
        # 跳过表格/图片标题
        if is_table_title(text) or is_figure_title(text):
            continue

        # 1. 清除文本中的项目符号字符
        if RE_BULLET_PREFIX.match(text):
            # 找到段落中第一个run，移除开头的项目符号
            all_runs = para._element.findall('.//' + qn('w:r'))
            if all_runs:
                first_run = all_runs[0]
                t_elem = first_run.find(qn('w:t'))
                if t_elem is not None and t_elem.text:
                    new_text = RE_BULLET_PREFIX.sub('', t_elem.text, count=1)
                    if new_text != t_elem.text:
                        t_elem.text = new_text
                        # 保持空格属性
                        t_elem.set(qn('xml:space'), 'preserve')
                        cleaned += 1

        # 2. 清除Word列表编号属性(numPr)
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                pPr.remove(numPr)
                cleaned += 1

    if cleaned > 0:
        print(f"  预处理: 清除了 {cleaned} 个项目符号标记")

    return cleaned


# =============================================================================
# 段落格式设置
# =============================================================================

def set_paragraph_format(para, alignment=None, first_indent=None,
                         line_spacing=None, space_before=None, space_after=None,
                         keep_with_next=False):
    """设置段落格式"""
    pf = para.paragraph_format

    if alignment is not None:
        pf.alignment = alignment
    if first_indent is not None:
        pf.first_line_indent = first_indent
    if line_spacing is not None:
        pf.line_spacing = line_spacing
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    pf.keep_with_next = keep_with_next

    # 启用东亚换行规则（避头尾）：
    # - kinsoku: 禁止标点出现在行首/行尾的不当位置（如"。"不能在行首）
    # - overflowPunct: 允许行尾标点略微溢出页边距，避免被挤到下一行
    # - autoSpaceDE/DN: 中英文、中文与数字之间自动添加间距
    pPr = para._element.get_or_add_pPr()
    for tag, val in [('w:kinsoku', '1'), ('w:overflowPunct', '1'),
                     ('w:autoSpaceDE', '1'), ('w:autoSpaceDN', '1')]:
        elem = pPr.find(qn(tag))
        if elem is None:
            elem = parse_xml(f'<{tag} {nsdecls("w")} w:val="{val}"/>')
            _insert_in_order(pPr, elem, _PPR_ORDER)
        else:
            elem.set(qn('w:val'), val)


def format_body_paragraph(para):
    """格式化正文段落"""
    set_paragraph_font(para, FONT_FANGSONG, FONT_TNR, BODY_FONT_SIZE)
    set_paragraph_format(
        para,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_indent=BODY_FIRST_LINE_INDENT,
        line_spacing=BODY_LINE_SPACING,
        space_before=BODY_PARA_BEFORE,
        space_after=BODY_PARA_AFTER,
    )


def set_outline_level(para, outline_level):
    """设置段落的大纲级别（用于目录生成）。

    outline_level: 0=一级, 1=二级, 2=三级, etc.
    对应 Word 中的"大纲级别"属性，TOC域根据此属性提取目录条目。
    """
    pPr = para._element.get_or_add_pPr()
    oLvl = pPr.find(qn('w:outlineLvl'))
    if oLvl is None:
        oLvl = parse_xml(f'<w:outlineLvl {nsdecls("w")} w:val="{outline_level}"/>')
        _insert_in_order(pPr, oLvl, _PPR_ORDER)
    else:
        oLvl.set(qn('w:val'), str(outline_level))


# 目录书签计数器（全局，每次排版重置）
_toc_bookmark_id = 0

def add_toc_bookmark(para, bm_name):
    """在标题段落中添加一个命名的书签，供目录PAGEREF域引用。

    书签以 bookmarkStart + bookmarkEnd 对的形式插入在段落内部，
    使 PAGEREF 域更新时能正确定位到该段落所在的页。
    """
    global _toc_bookmark_id
    _toc_bookmark_id += 1
    bm_id = _toc_bookmark_id

    p_elem = para._element

    # bookmarkStart 插入到段落第一个子元素前
    bm_start = parse_xml(
        f'<w:bookmarkStart {nsdecls("w")} w:id="{bm_id}" w:name="{bm_name}"/>'
    )
    p_elem.insert(0, bm_start)

    # bookmarkEnd 追加到段落末尾
    bm_end = parse_xml(
        f'<w:bookmarkEnd {nsdecls("w")} w:id="{bm_id}"/>'
    )
    p_elem.append(bm_end)


def format_heading(para, level):
    """格式化标题段落

    v2.7.3 变更：
    - 一级和二级标题首行缩进2个字符（与正文一致）
    - 一级和二级标题设置大纲级别（0和1），以便生成目录
    - 一级和二级标题字体强制黑色（RGB 000000），避免蓝色
    """
    FONT_COLOR_BLACK = RGBColor(0, 0, 0)

    if level == 1:
        set_paragraph_font(para, H1_FONT, FONT_TNR, H1_SIZE, H1_BOLD,
                           color=FONT_COLOR_BLACK)
        set_paragraph_format(
            para,
            first_indent=BODY_FIRST_LINE_INDENT,  # v2.7.3: 首行缩进2字符
            line_spacing=BODY_LINE_SPACING,
            space_before=BODY_PARA_BEFORE,
            space_after=BODY_PARA_AFTER,
            keep_with_next=True,
        )
        set_outline_level(para, 0)  # v2.7.3: 大纲级别0（一级）
    elif level == 2:
        set_paragraph_font(para, H2_FONT, FONT_TNR, H2_SIZE, H2_BOLD,
                           color=FONT_COLOR_BLACK)
        set_paragraph_format(
            para,
            first_indent=BODY_FIRST_LINE_INDENT,  # v2.7.3: 首行缩进2字符
            line_spacing=BODY_LINE_SPACING,
            space_before=BODY_PARA_BEFORE,
            space_after=BODY_PARA_AFTER,
            keep_with_next=True,
        )
        set_outline_level(para, 1)  # v2.7.3: 大纲级别1（二级）
    else:
        # 三级及以下：字体字号与正文一致
        set_paragraph_font(para, H3_FONT, FONT_TNR, H3_SIZE, H3_BOLD)
        set_paragraph_format(
            para,
            line_spacing=BODY_LINE_SPACING,
            space_before=BODY_PARA_BEFORE,
            space_after=BODY_PARA_AFTER,
            keep_with_next=True,
        )
        para.paragraph_format.first_line_indent = Cm(0)


def format_table_title_paragraph(para):
    """格式化表格标题段落（如 '表3 XXX'）"""
    set_paragraph_font(para, TABLE_TITLE_FONT, FONT_TNR, TABLE_TITLE_SIZE, TABLE_TITLE_BOLD)
    set_paragraph_format(
        para,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_indent=Cm(0),
        line_spacing=BODY_LINE_SPACING,
        keep_with_next=True,  # 表格标题不与表格分跨两页
    )


def format_figure_title_paragraph(para):
    """格式化图例标题段落"""
    set_paragraph_font(para, FIGURE_TITLE_FONT, FONT_TNR, FIGURE_TITLE_SIZE)
    set_paragraph_format(
        para,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=Cm(0),
        line_spacing=BODY_LINE_SPACING,
    )


# =============================================================================
# 目录插入
# =============================================================================

def _create_toc_entry_element(level, heading_text, bookmark_name="_Toc"):
    """创建一个目录条目段落XML元素。

    Args:
        level: 标题层级（1=一级, 2=二级）
        heading_text: 标题文本
        bookmark_name: 对应标题的书签名称，用于PAGEREF域引用

    Returns:
        lxml.etree.Element: w:p 元素
    """
    # 目录条目格式：
    # 一级：黑体四号加粗，行距1.5倍，段后6磅
    # 二级：楷体小四加黑，行距1.5倍
    # 两者均首行缩进2字符
    # 注意：不使用 w:pStyle 引用 TOC 样式（文档中可能不存在该样式），
    # 所有格式通过显式 pPr/rPr 设置
    if level == 1:
        cn_font = TOC_L1_FONT
        size = TOC_L1_SIZE
        bold_tag = '<w:b/><w:bCs/>'  # 一级加粗
        space_after_val = '120'  # 6磅 = 120 twentieths-of-a-point
    elif level == 2:
        cn_font = TOC_L2_FONT
        size = TOC_L2_SIZE
        bold_tag = '<w:b/><w:bCs/>'  # 二级加粗
        space_after_val = '0'
    else:
        cn_font = TOC_L3_FONT
        size = TOC_L3_SIZE
        bold_tag = ''  # 三级不加粗
        space_after_val = '0'

    # 行距1.5倍 = 360/240 (spacing line = 360, lineRule=auto where 240=single)
    half_points = str(int(size.pt * 2))

    # 首行缩进2字符：同时设 firstLineChars 和 firstLine（twips）
    # firstLine 是绝对值备份，确保不含 firstLineChars 属性的阅读器也能正确缩进
    # 0.85cm ≈ 481 twips
    first_line_twips = '481'

    # XML 中的 heading_text 需要 XML 转义
    escaped_text = heading_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;')

    p_xml = (
        f'<w:p {nsdecls("w")}>'
        f'<w:pPr>'
        f'<w:tabs>'
        f'<w:tab w:val="right" w:leader="dot" w:pos="8296"/>'
        f'</w:tabs>'
        f'<w:spacing w:after="{space_after_val}" w:line="360" w:lineRule="auto"/>'
        f'<w:ind w:firstLineChars="200" w:firstLine="{first_line_twips}"/>'
        f'<w:rPr>'
        f'<w:rFonts w:eastAsia="{cn_font}" w:ascii="{FONT_TNR}" w:hAnsi="{cn_font}"/>'
        f'<w:sz w:val="{half_points}"/>'
        f'<w:szCs w:val="{half_points}"/>'
        f'{bold_tag}'
        f'<w:color w:val="000000"/>'
        f'</w:rPr>'
        f'</w:pPr>'
        f'<w:r>'
        f'<w:rPr>'
        f'<w:rFonts w:eastAsia="{cn_font}" w:ascii="{FONT_TNR}" w:hAnsi="{cn_font}"/>'
        f'<w:sz w:val="{half_points}"/>'
        f'<w:szCs w:val="{half_points}"/>'
        f'{bold_tag}'
        f'<w:color w:val="000000"/>'
        f'</w:rPr>'
        f'<w:t xml:space="preserve">{escaped_text}</w:t>'
        f'</w:r>'
        f'<w:r><w:tab/></w:r>'
        f'<w:r>'
        f'<w:rPr>'
        f'<w:rFonts w:eastAsia="{cn_font}" w:ascii="{FONT_TNR}" w:hAnsi="{cn_font}"/>'
        f'<w:sz w:val="{half_points}"/>'
        f'<w:szCs w:val="{half_points}"/>'
        f'<w:color w:val="000000"/>'
        f'</w:rPr>'
        f'<w:fldChar w:fldCharType="begin"/>'
        f'</w:r>'
        f'<w:r>'
        f'<w:rPr>'
        f'<w:rFonts w:eastAsia="{cn_font}" w:ascii="{FONT_TNR}" w:hAnsi="{cn_font}"/>'
        f'<w:sz w:val="{half_points}"/>'
        f'<w:szCs w:val="{half_points}"/>'
        f'<w:color w:val="000000"/>'
        f'</w:rPr>'
        f'<w:instrText xml:space="preserve"> PAGEREF {bookmark_name} \\h </w:instrText>'
        f'</w:r>'
        f'<w:r>'
        f'<w:rPr>'
        f'<w:rFonts w:eastAsia="{cn_font}" w:ascii="{FONT_TNR}" w:hAnsi="{cn_font}"/>'
        f'<w:sz w:val="{half_points}"/>'
        f'<w:szCs w:val="{half_points}"/>'
        f'<w:color w:val="000000"/>'
        f'</w:rPr>'
        f'<w:fldChar w:fldCharType="separate"/>'
        f'</w:r>'
        f'<w:r>'
        f'<w:rPr>'
        f'<w:rFonts w:eastAsia="{cn_font}" w:ascii="{FONT_TNR}" w:hAnsi="{cn_font}"/>'
        f'<w:sz w:val="{half_points}"/>'
        f'<w:szCs w:val="{half_points}"/>'
        f'<w:color w:val="000000"/>'
        f'</w:rPr>'
        f'<w:t>1</w:t>'
        f'</w:r>'
        f'<w:r>'
        f'<w:rPr>'
        f'<w:rFonts w:eastAsia="{cn_font}" w:ascii="{FONT_TNR}" w:hAnsi="{cn_font}"/>'
        f'<w:sz w:val="{half_points}"/>'
        f'<w:szCs w:val="{half_points}"/>'
        f'<w:color w:val="000000"/>'
        f'</w:rPr>'
        f'<w:fldChar w:fldCharType="end"/>'
        f'</w:r>'
        f'</w:p>'
    )

    return parse_xml(p_xml)


def insert_toc_entries(doc, toc_title_idx, headings_list, cover_end_idx=None):
    """在目录标题后插入目录条目段落（默认显示到二级目录）。

    策略：
    1. 扫描目录标题与第一个H1之间的段落，如果全是空段落或仅含fldChar域代码，
       则清除它们并插入手动构建的目录条目
    2. 如果目录区域已有实际内容（非空、非域代码），则保留不动
    3. 手动条目包含标题文本+制表位+PAGEREF页码域，打开Word后可更新为准确页码

    Args:
        doc: Document 对象
        toc_title_idx: "目录"标题段落的索引
        headings_list: [(level, text, bookmark_name), ...] 所有一级和二级标题
        cover_end_idx: 封面结束段落索引
    """
    body = doc.element.body
    all_paras = list(body.findall(qn('w:p')))

    # 找到目录标题段落元素
    toc_title_elem = all_paras[toc_title_idx]

    # 确定目录区域的结束位置：目录标题后到第一个H1标题前
    # 从目录标题开始往后扫描，找到第一个有outlineLvl或被识别为H1的段落
    toc_area_end_idx = None
    for j in range(toc_title_idx + 1, len(all_paras)):
        p_elem = all_paras[j]
        pPr = p_elem.find(qn('w:pPr'))
        # 检查是否有大纲级别设置（H1/H2标题会有）
        if pPr is not None:
            oLvl = pPr.find(qn('w:outlineLvl'))
            if oLvl is not None and oLvl.get(qn('w:val')) in ('0', '1'):
                toc_area_end_idx = j
                break
        # 也检查段落文本是否匹配H1模式
        text_elems = p_elem.findall('.//' + qn('w:t'))
        para_text = ''.join(t.text for t in text_elems if t.text).strip()
        if RE_H1.match(para_text):
            toc_area_end_idx = j
            break

    # 如果没找到H1，目录区域延伸到文档末尾之前
    if toc_area_end_idx is None:
        return  # 无法确定目录区域，不插入

    # 检查目录区域现有内容：
    # - 有文本内容的段落（无论是否含fldChar）视为真实内容 → 保留不动
    # - 空段落（无文本内容）视为可清除 → 标记移除
    has_real_content = False
    paras_to_remove = []
    for j in range(toc_title_idx + 1, toc_area_end_idx):
        p_elem = all_paras[j]
        text_elems = p_elem.findall('.//' + qn('w:t'))
        para_text = ''.join(t.text for t in text_elems if t.text).strip()

        if para_text:
            # 有文本就是真实内容，无论是手动条目还是Word生成的TOC域
            has_real_content = True
            break

        # 空段落，可以清除
        paras_to_remove.append(p_elem)

    # 如果已有实际内容（非空非域代码），保留不动
    if has_real_content:
        return

    # 清除目录区域的空段落和域代码段落
    for p_elem in paras_to_remove:
        body.remove(p_elem)

    # 过滤标题：只取一级和二级
    toc_headings = [(lv, txt, bm) for lv, txt, bm in headings_list if lv <= 2]

    if not toc_headings:
        return  # 没有标题，无需插入

    # 在目录标题后插入条目段落
    last_elem = toc_title_elem
    for level, heading_text, bookmark_name in toc_headings:
        entry_elem = _create_toc_entry_element(level, heading_text, bookmark_name)
        last_elem.addnext(entry_elem)
        last_elem = entry_elem

    # 在所有条目后插入一行提示
    note_xml = (
        f'<w:p {nsdecls("w")}>'
        f'<w:pPr>'
        f'<w:spacing w:line="360" w:lineRule="auto"/>'
        f'<w:rPr>'
        f'<w:rFonts w:eastAsia="{FONT_FANGSONG}" w:ascii="{FONT_TNR}" w:hAnsi="{FONT_FANGSONG}"/>'
        f'<w:sz w:val="24"/>'  # 小四12pt
        f'<w:szCs w:val="24"/>'
        f'<w:color w:val="808080"/>'  # 灰色
        f'</w:rPr>'
        f'</w:pPr>'
        f'<w:r>'
        f'<w:rPr>'
        f'<w:rFonts w:eastAsia="{FONT_FANGSONG}" w:ascii="{FONT_TNR}" w:hAnsi="{FONT_FANGSONG}"/>'
        f'<w:sz w:val="24"/>'
        f'<w:szCs w:val="24"/>'
        f'<w:color w:val="808080"/>'
        f'</w:rPr>'
        f'<w:t>（请在Word中右键此目录 → 更新域 → 更新整个目录，以获取准确页码）</w:t>'
        f'</w:r>'
        f'</w:p>'
    )
    note_elem = parse_xml(note_xml)
    last_elem.addnext(note_elem)


# =============================================================================
# 封面格式设置
# =============================================================================

def detect_cover_paragraphs(doc):
    """
    检测封面区域的段落索引范围。
    封面结束于第一个分页符（page break）之前。
    返回 (start_idx, end_idx) 表示封面段落的索引范围（含两端），
    如果没有检测到分页符则返回 None。

    v2.7.3: 若文档无分页符但调用方声明 has_cover=True，
    回退策略——找到第一个一级标题之前的非空段落作为封面末尾。
    典型场景：用户文档封面和正文之间没有分页符，只有空行分隔。
    """
    for i, para in enumerate(doc.paragraphs):
        # 检查段落内是否包含分页符
        for run in para.runs:
            for child in run._element:
                if child.tag == qn('w:br'):
                    br_type = child.get(qn('w:type'))
                    if br_type == 'page':
                        # 分页符之前的段落都是封面
                        return (0, i)
        # 也检查段落本身是否是 page break before
        if para.paragraph_format.page_break_before:
            return (0, max(0, i - 1))
    return None


def _infer_cover_range(doc):
    """v2.7.3: 当 detect_cover_paragraphs 返回 None 但 has_cover=True 时，
    推断封面范围。

    策略：从文档开头向下扫描，优先把"目录"作为封面结束边界；
    如果没有目录，再找到第一个被识别为一级标题的段落，
    该边界之前的所有段落（含末尾空行）视为封面。
    如果没有找到目录或一级标题，返回 None。
    """
    boundary_idx = None
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        if text in ('目录', '目 录'):
            boundary_idx = i
            break
        level = detect_heading_level(text, para)
        if level == 1:
            boundary_idx = i
            break

    if boundary_idx is None:
        return None

    # 封面末尾 = 目录或一级标题前的最后一个非空段落
    # 如果边界前有空行，封面到空行前一行为止
    cover_end = boundary_idx - 1
    while cover_end >= 0 and not doc.paragraphs[cover_end].text.strip():
        cover_end -= 1

    if cover_end < 0:
        return None

    return (0, cover_end)


def format_cover_page(doc, cover_range):
    """
    格式化封面页：
    - 标题：黑体二号，加粗，居中，位于页面上方约1/3处
    - 副标题/说明文字：楷体四号，居中（标题与署名之间的所有非空段落）
    - 署名和日期：楷体四号，居中，位于页面底部

    封面段落的识别规则：
    - 第一个非空段落 = 标题
    - 末尾连续的"含公司/年/月"段落 = 署名/日期
    - 标题与署名之间的所有非空段落 = 副标题（统一居中处理）

    v2.7.2 修复：
    - 副标题识别不再要求以括号开头，所有"标题-署名之间的非空段落"均居中
    - 落款 space_before 改为动态计算，避免封面副标题较多时落款溢出第二页
    """
    start_idx, end_idx = cover_range
    paragraphs = doc.paragraphs[start_idx:end_idx + 1]

    # 将非空段落分类
    non_empty = [(i, p) for i, p in enumerate(paragraphs) if p.text.strip()]

    if not non_empty:
        return

    # 找到署名和日期段落（通常是最后的2-3个非空段落）
    # 署名特征：含"公司""有限""咨询""研究院"等机构后缀；日期特征：含"年""月"
    # v2.7.2：排除含冒号的段落（如"编制日期：2026 年 5 月"），此类应归副标题
    #         机构识别去掉过于宽泛的"管理""集团""编制"独立匹配，
    #         改为要求与机构后缀连用（如"管理咨询""集团有限公司"）
    author_date_indices = []
    org_pattern = re.compile(
        r'(公司|有限|咨询|研究院|研究所|集团有限|管理咨询|事务所|中心$|学院$|大学$)'
    )
    pure_date_pattern = re.compile(r'^\D*\d{4}\s*年\s*\d{1,2}\s*月\s*\d{0,2}\s*日?\s*$')
    cn_date_pattern = re.compile(r'^[二三四五六七八九〇零一]+年[^：:]{0,10}月?$')

    for idx in range(len(non_empty) - 1, -1, -1):
        _, p = non_empty[idx]
        text = p.text.strip()
        # 含冒号（半/全角）一律不算落款——多半是"编制日期：…""提交对象：…"等副标题
        if '：' in text or ':' in text:
            break
        # 匹配三种"纯落款"形态：机构名 / 纯阿拉伯数字日期 / 纯中文数字日期
        if (org_pattern.search(text) or
                pure_date_pattern.match(text) or
                cn_date_pattern.match(text)):
            author_date_indices.insert(0, idx)
        else:
            break  # 遇到非署名/日期内容就停止

    title_idx = 0  # 第一个非空段落是标题

    # v2.7.2：副标题 = 标题与署名之间的所有非空段落（不再要求括号开头）
    # 这样 "以XX为主题的研究"、"(讨论稿)"、"内部使用版" 等所有副标题均能被识别并居中
    subtitle_indices = [
        idx for idx in range(len(non_empty))
        if idx != title_idx and idx not in author_date_indices
    ]

    # ---- 格式化标题 ----
    _, title_para = non_empty[title_idx]
    set_paragraph_font(title_para, COVER_TITLE_FONT, FONT_TNR,
                       COVER_TITLE_SIZE, COVER_TITLE_BOLD)
    # 强制封面标题为黑色（清除可能残留的蓝色超链接等颜色）
    for r in title_para._element.findall('.//' + qn('w:r')):
        _set_run_element_font(r, COVER_TITLE_FONT, FONT_TNR,
                              COVER_TITLE_SIZE, COVER_TITLE_BOLD, color=RGBColor(0, 0, 0))
    title_space_before = Pt(80)
    title_space_after = Pt(20)
    set_paragraph_format(
        title_para,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=Cm(0),
        line_spacing=BODY_LINE_SPACING,
        space_before=title_space_before,
        space_after=title_space_after,
    )

    # ---- 格式化副标题/说明文字（统一居中）----
    subtitle_space_before = Pt(6)
    subtitle_space_after = Pt(6)
    for idx in subtitle_indices:
        _, sub_para = non_empty[idx]
        set_paragraph_font(sub_para, COVER_SUBTITLE_FONT, FONT_TNR,
                           COVER_SUBTITLE_SIZE)
        set_paragraph_format(
            sub_para,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_indent=Cm(0),
            line_spacing=BODY_LINE_SPACING,
            space_before=subtitle_space_before,
            space_after=subtitle_space_after,
        )

    # ---- 格式化署名和日期：动态推算 space_before 推到页面底部 ----
    if author_date_indices:
        # v2.7.2：动态计算 space_before，防止封面副标题过多时落款溢出第二页
        #
        # 估算各区块占用高度（单位：pt）：
        #   版心高度 = 225mm ≈ 638pt
        #   标题区 = 段前 + (字号 × 行距) + 段后
        #   每个副标题 = 段前 + (字号 × 行距) + 段后
        #   每个署名/日期 = 段前 + (字号 × 行距) + 段后
        page_height_pt = 225.0 / 25.4 * 72.0  # ≈ 638pt

        title_height = (title_space_before.pt
                        + COVER_TITLE_SIZE.pt * BODY_LINE_SPACING
                        + title_space_after.pt)

        subtitle_unit_height = (subtitle_space_before.pt
                                + COVER_SUBTITLE_SIZE.pt * BODY_LINE_SPACING
                                + subtitle_space_after.pt)
        subtitle_total_height = subtitle_unit_height * len(subtitle_indices)

        # 落款区高度：每段 = (字号 × 行距)；段前段后均为 0
        author_unit_height = COVER_AUTHOR_SIZE.pt * BODY_LINE_SPACING
        author_total_height = author_unit_height * len(author_date_indices)

        # 已用高度 + 缓冲（80pt 安全余量）
        used_height = title_height + subtitle_total_height
        safety_buffer = 80.0
        available_for_space_before = (page_height_pt
                                      - used_height
                                      - author_total_height
                                      - safety_buffer)

        # 上下限兜底：[100pt, 420pt]
        # - 下限 100pt：保证署名与副标题之间有清晰视觉分隔
        # - 上限 420pt：避免极端情况下间距过大导致渲染异常
        space_before_pt = max(100.0, min(420.0, available_for_space_before))

        first_author_real_idx = author_date_indices[0]
        _, first_author_para = non_empty[first_author_real_idx]

        set_paragraph_font(first_author_para, COVER_AUTHOR_FONT, FONT_TNR,
                           COVER_AUTHOR_SIZE)
        set_paragraph_format(
            first_author_para,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_indent=Cm(0),
            line_spacing=BODY_LINE_SPACING,
            space_before=Pt(space_before_pt),
            space_after=Pt(0),
        )
        # 显式禁止在落款前分页（防御性）
        first_author_para.paragraph_format.page_break_before = False

        # 后续的署名/日期段落正常间距
        for idx in author_date_indices[1:]:
            _, p = non_empty[idx]
            set_paragraph_font(p, COVER_AUTHOR_FONT, FONT_TNR,
                               COVER_AUTHOR_SIZE)
            set_paragraph_format(
                p,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                first_indent=Cm(0),
                line_spacing=BODY_LINE_SPACING,
                space_before=Pt(0),
                space_after=Pt(0),
            )

    # 处理封面中的空段落（设为最小间距，不影响布局）
    for p in paragraphs:
        if not p.text.strip():
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            # 设置字号尽可能小以减少空段落占用空间
            for run in p.runs:
                run.font.size = Pt(2)


# =============================================================================
# 表格格式设置
# =============================================================================

def _estimate_table_width_need(table):
    """估算表格所需的宽度级别。

    Returns:
        'normal': 可在纵向页面容纳
        'wide':   需要横向页面
    """
    num_cols = max((len(row.cells) for row in table.rows), default=0)

    # 规则1：列数 >= 6 通常需要横向
    if num_cols >= 6:
        return 'wide'

    # 规则2：检查表格内容的平均宽度
    # 计算所有单元格文本的总长度
    total_text_len = 0
    max_cell_len = 0
    for row in table.rows:
        row_text_len = 0
        for cell in row.cells:
            cell_len = len(cell.text.strip())
            row_text_len += cell_len
            max_cell_len = max(max_cell_len, cell_len)
        total_text_len = max(total_text_len, row_text_len)

    # 纵向可用宽度约156mm，小五号字(9pt)每个中文字约3.2mm，约可容纳48个字符
    # 考虑列间距，每行总字符超过50且列数>=5时需要横向
    if total_text_len > 50 and num_cols >= 5:
        return 'wide'

    # 规则3：单个单元格文本特别长（>60字符）且列数>=4
    if max_cell_len > 60 and num_cols >= 4:
        return 'wide'

    return 'normal'


def _set_table_autofit(table):
    """设置表格自动适应页面宽度（不超出页面边界）。

    设置 tblW type="pct" val="5000"，表示占页面可用宽度的100%。
    同时设置 tblLayout 为 autofit，让 Word 自动调整列宽。
    """
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tbl_pr)

    # 设置表格总宽度为页面100%
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = parse_xml(f'<w:tblW {nsdecls("w")} w:type="pct" w:w="5000"/>')
        tbl_pr.insert(0, tbl_w)
    else:
        tbl_w.set(qn('w:type'), 'pct')
        tbl_w.set(qn('w:w'), '5000')

    # 设置自动布局（让Word自动分配列宽）
    layout = tbl_pr.find(qn('w:tblLayout'))
    if layout is not None:
        tbl_pr.remove(layout)
    # 不设置 tblLayout 即默认 autofit


def format_table(table):
    """
    格式化表格为三线表风格:
    - 顶线和底线较粗（1.5pt），栏目线稍细（0.75pt）
    - 表头行：楷体_GB2312，小五，加粗
    - 数据行：宋体，小五
    - 居中对齐
    - 自动适应页面宽度
    """
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_autofit(table)

    # 设置三线表边框
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')

    # 清除所有边框，然后只设置三线
    borders_xml = f'''<w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>
        <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>
        <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    </w:tblBorders>'''
    borders_elem = parse_xml(borders_xml)

    existing_borders = tbl_pr.find(qn('w:tblBorders'))
    if existing_borders is not None:
        tbl_pr.remove(existing_borders)
    tbl_pr.append(borders_elem)

    for i, row in enumerate(table.rows):
        for cell in row.cells:
            # 设置单元格内段落格式
            for para in cell.paragraphs:
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)

                if i == 0:
                    # 表头行
                    set_paragraph_font(para, TABLE_HEADER_FONT, FONT_TNR,
                                       TABLE_HEADER_SIZE, TABLE_HEADER_BOLD)
                else:
                    # 数据行
                    set_paragraph_font(para, TABLE_BODY_FONT, FONT_TNR,
                                       TABLE_BODY_SIZE)

            # 为表头行添加下边框（栏目线）
            if i == 0:
                tc_pr = cell._tc.get_or_add_tcPr()
                cell_borders_xml = f'''<w:tcBorders {nsdecls("w")}>
                    <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>
                </w:tcBorders>'''
                cell_borders = parse_xml(cell_borders_xml)
                existing = tc_pr.find(qn('w:tcBorders'))
                if existing is not None:
                    tc_pr.remove(existing)
                tc_pr.append(cell_borders)


def _find_table_title_para(body, tbl_elem):
    """从表格元素向前查找表格标题段落。

    跳过空段落和 sectPr 空段落，找到第一个有文本内容的段落。
    如果该段落文本匹配表格标题格式（如 '表1 XXX'），返回该段落元素；
    否则返回 None。
    """
    prev = tbl_elem.getprevious()
    while prev is not None:
        tag = prev.tag.split('}')[-1] if '}' in prev.tag else prev.tag
        if tag == 'p':
            texts = prev.findall('.//' + qn('w:t'))
            text = ''.join(t.text or '' for t in texts).strip()
            if text:
                if is_table_title(text):
                    return prev
                else:
                    return None
        prev = prev.getprevious()
    return None


def _wrap_wide_tables_landscape(doc, wide_table_indices):
    """为宽表格插入横向分节符。

    v2.7.5 重写：修正 OOXML sectPr 语义。
    OOXML 规范：段落 pPr 中的 sectPr 定义的是**到该段落为止**的节的属性。
    因此 landscape sectPr 必须放在表格**之后**的段落中，才能使表格标题+表格
    都包含在 landscape 节内。

    分节结构示例：
    [正文段落(pPr: portrait-nextPage)] → [表格标题] → <w:tbl>
    → [空段落(pPr: landscape-nextPage)] → [后续正文]

    解释：
    - portrait-nextPage 在正文段落中：定义从上一个 sectPr 到该段落的节为
      portrait，nextPage 使后续内容另起一页（landscape 节起始新页）。
    - landscape-nextPage 在表格后空段落中：定义从表格标题到该空段落的节为
      landscape，nextPage 使后续内容另起一页（恢复 portrait）。
    - 后续内容继承 body 默认 sectPr（portrait）。
    """
    def mm_to_twips(mm):
        return int(mm / 25.4 * 1440)

    pg_w_portrait = mm_to_twips(210)    # A4纵向宽度
    pg_h_portrait = mm_to_twips(297)    # A4纵向高度
    mar_top = mm_to_twips(37)           # 天头
    mar_bottom = mm_to_twips(35)        # 下边距
    mar_left = mm_to_twips(28)          # 订口
    mar_right = mm_to_twips(26)         # 右边距
    mar_footer = mm_to_twips(7)         # 页脚距

    for idx in reversed(wide_table_indices):
        # 倒序处理，避免索引偏移
        if idx >= len(doc.tables):
            continue
        tbl_elem = doc.tables[idx]._tbl

        # 找到表格标题段落
        title_para = _find_table_title_para(doc.element.body, tbl_elem)

        # --- 1. 在表格标题前的段落中插入 portrait nextPage sectPr ---
        # 作用：定义前面的节为 portrait，nextPage 使 landscape 节另起新页
        portrait_anchor_xml = (
            f'<w:sectPr {nsdecls("w")}>'
            f'  <w:pgSz w:w="{pg_w_portrait}" w:h="{pg_h_portrait}" w:orient="portrait"/>'
            f'  <w:pgMar w:top="{mar_top}" w:right="{mar_right}"'
            f'          w:bottom="{mar_bottom}" w:left="{mar_left}"'
            f'          w:header="720" w:footer="{mar_footer}"/>'
            f'  <w:sectType w:val="nextPage"/>'
            f'</w:sectPr>'
        )
        portrait_anchor = parse_xml(portrait_anchor_xml)

        if title_para is not None:
            # 找到表格标题前的一个段落
            prev_elem = title_para.getprevious()
            while prev_elem is not None:
                tag = prev_elem.tag.split('}')[-1] if '}' in prev_elem.tag else prev_elem.tag
                if tag == 'p':
                    break
                prev_elem = prev_elem.getprevious()

            if prev_elem is not None and prev_elem.tag.endswith('}p'):
                prev_pPr = prev_elem.find(qn('w:pPr'))
                if prev_pPr is None:
                    prev_pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
                    prev_elem.insert(0, prev_pPr)
                _insert_in_order(prev_pPr, portrait_anchor, _PPR_ORDER)
            else:
                # 前面没有段落，创建一个空段落
                pre_para = parse_xml(
                    f'<w:p {nsdecls("w")}><w:pPr></w:pPr></w:p>'
                )
                title_para.addprevious(pre_para)
                pre_pPr = pre_para.find(qn('w:pPr'))
                _insert_in_order(pre_pPr, portrait_anchor, _PPR_ORDER)
        else:
            # 回退：找不到表格标题，在表格前插入空段落承载 portrait sectPr
            pre_para_xml = (
                f'<w:p {nsdecls("w")}>'
                f'  <w:pPr>'
                f'    <w:sectPr>'
                f'      <w:pgSz w:w="{pg_w_portrait}" w:h="{pg_h_portrait}" w:orient="portrait"/>'
                f'      <w:pgMar w:top="{mar_top}" w:right="{mar_right}"'
                f'              w:bottom="{mar_bottom}" w:left="{mar_left}"'
                f'              w:header="720" w:footer="{mar_footer}"/>'
                f'      <w:sectType w:val="nextPage"/>'
                f'    </w:sectPr>'
                f'  </w:pPr>'
                f'</w:p>'
            )
            pre_para = parse_xml(pre_para_xml)
            tbl_elem.addprevious(pre_para)

        # --- 2. 在表格后插入 landscape nextPage 空段落 ---
        # 关键：landscape sectPr 放在表格之后，使"表格标题+表格"都在 landscape 节中
        landscape_post_xml = (
            f'<w:p {nsdecls("w")}>'
            f'  <w:pPr>'
            f'    <w:sectPr>'
            f'      <w:pgSz w:w="{pg_h_portrait}" w:h="{pg_w_portrait}" w:orient="landscape"/>'
            f'      <w:pgMar w:top="{mar_left}" w:right="{mar_top}"'
            f'              w:bottom="{mar_right}" w:left="{mar_bottom}"'
            f'              w:header="720" w:footer="{mar_footer}"/>'
            f'      <w:sectType w:val="nextPage"/>'
            f'    </w:sectPr>'
            f'  </w:pPr>'
            f'</w:p>'
        )
        post_landscape_para = parse_xml(landscape_post_xml)
        tbl_elem.addnext(post_landscape_para)


# =============================================================================
# 页码设置
# =============================================================================

def setup_page_numbers(doc, style='dash', has_cover=False):
    """
    设置页码：
    - 5号半角宋体阿拉伯数字
    - 一字线格式 "— X —"（style='dash'）或纯数字（style='plain'）
    - 奇数页居右空一字，偶数页居左空一字
    """
    for section in doc.sections:
        # 启用奇偶页不同
        section.different_first_page_header_footer = has_cover
        sectPr = section._sectPr

        # 设置奇偶页不同
        even_odd = sectPr.find(qn('w:evenAndOddHeaders'))
        if even_odd is None:
            even_odd = parse_xml(f'<w:evenAndOddHeaders {nsdecls("w")}/>')
            sectPr.append(even_odd)

        # 页码距版心下边缘7mm
        # w:footer 属性使用 twips 单位（1/20点 = 1/1440英寸）
        # 7mm = 7/25.4 * 1440 ≈ 397 twips
        footer_twips = int(7 / 25.4 * 1440)
        pgMar = sectPr.find(qn('w:pgMar'))
        if pgMar is not None:
            pgMar.set(qn('w:footer'), str(footer_twips))

        # 奇数页页脚（默认页脚 = 奇数页）
        footer_odd = section.footer
        footer_odd.is_linked_to_previous = False
        for p in footer_odd.paragraphs:
            p.clear()

        if footer_odd.paragraphs:
            p_odd = footer_odd.paragraphs[0]
        else:
            p_odd = footer_odd.add_paragraph()

        p_odd.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if style == 'dash':
            _add_page_number_with_dash(p_odd)
        else:
            _add_plain_page_number(p_odd)

        # 偶数页页脚
        footer_even = section.even_page_footer
        footer_even.is_linked_to_previous = False
        for p in footer_even.paragraphs:
            p.clear()

        if footer_even.paragraphs:
            p_even = footer_even.paragraphs[0]
        else:
            p_even = footer_even.add_paragraph()

        p_even.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if style == 'dash':
            _add_page_number_with_dash(p_even)
        else:
            _add_plain_page_number(p_even)


def _add_page_number_with_dash(paragraph):
    """添加 '— X —' 格式的页码"""
    run1 = paragraph.add_run('— ')
    set_run_font(run1, FONT_SONGTI, FONT_TNR, PAGE_NUM_SIZE)

    # 插入PAGE字段
    fld_xml = (
        f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE \\* MERGEFORMAT ">'
        f'<w:r><w:rPr><w:rFonts w:ascii="{FONT_SONGTI}" w:eastAsia="{FONT_SONGTI}"/>'
        f'<w:sz w:val="{int(PAGE_NUM_SIZE.pt * 2)}"/></w:rPr>'
        f'<w:t>1</w:t></w:r></w:fldSimple>'
    )
    fld_elem = parse_xml(fld_xml)
    paragraph._element.append(fld_elem)

    run2 = paragraph.add_run(' —')
    set_run_font(run2, FONT_SONGTI, FONT_TNR, PAGE_NUM_SIZE)


def _add_plain_page_number(paragraph):
    """添加纯数字页码"""
    fld_xml = (
        f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE \\* MERGEFORMAT ">'
        f'<w:r><w:rPr><w:rFonts w:ascii="{FONT_SONGTI}" w:eastAsia="{FONT_SONGTI}"/>'
        f'<w:sz w:val="{int(PAGE_NUM_SIZE.pt * 2)}"/></w:rPr>'
        f'<w:t>1</w:t></w:r></w:fldSimple>'
    )
    fld_elem = parse_xml(fld_xml)
    paragraph._element.append(fld_elem)


# =============================================================================
# 页眉设置
# =============================================================================

def setup_header(doc, left_text='', right_text=''):
    """
    设置页眉：
    - 左侧：报告名称，楷体五号加黑
    - 右侧：编制单位名称，黑体小五
    - 细横线分隔
    """
    if not left_text and not right_text:
        return

    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False

        # 清空现有内容
        for p in header.paragraphs:
            p.clear()

        if header.paragraphs:
            p = header.paragraphs[0]
        else:
            p = header.add_paragraph()

        # 左侧文本
        if left_text:
            run_left = p.add_run(left_text)
            set_run_font(run_left, HEADER_LEFT_FONT, FONT_TNR,
                         HEADER_LEFT_SIZE, HEADER_LEFT_BOLD)

        # Tab + 右侧文本
        if right_text:
            run_tab = p.add_run('\t')
            run_right = p.add_run(right_text)
            set_run_font(run_right, HEADER_RIGHT_FONT, FONT_TNR,
                         HEADER_RIGHT_SIZE)

        # 设置Tab停止位（右对齐到版心右边缘）
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 添加下边框（细横线）
        pPr = p._element.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="000000"/>'
            f'</w:pBdr>'
        )
        existing_bdr = pPr.find(qn('w:pBdr'))
        if existing_bdr is not None:
            pPr.remove(existing_bdr)
        pPr.append(pBdr)




# =============================================================================
# v2.7：中文标点国标自动修正（GB/T 15834—2011）
# =============================================================================

def fix_punctuation_in_doc(doc):
    """遍历文档所有段落与表格单元格的 run，按 GB/T 15834 把中文上下文里的
    半角标点（逗号、句号、冒号、分号、问号、叹号、ASCII 直引号、半角括号、
    尖括号冒充书名号、--、... 等）修正为全角。

    跳过白名单：URL、邮箱、文件路径、版本号、千分位、小数点、行内代码。

    返回修正的字符数。
    """
    total_changes = 0

    def _fix_run_text(r_elem):
        nonlocal total_changes
        t = r_elem.find(qn('w:t'))
        if t is None or t.text is None:
            return
        new_text, changes = fix_chinese_punctuation(t.text)
        if changes > 0:
            t.text = new_text
            # 保留前后空格
            t.set(qn('xml:space'), 'preserve')
            total_changes += changes

    # 段落
    for para in doc.paragraphs:
        for r_elem in para._element.findall('.//' + qn('w:r')):
            _fix_run_text(r_elem)

    # 表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for r_elem in para._element.findall('.//' + qn('w:r')):
                        _fix_run_text(r_elem)

    # 页眉/页脚（保险起见，但实际页眉里几乎不会有 AI 半角问题）
    for section in doc.sections:
        for hf in (section.header, section.footer,
                   getattr(section, 'even_page_header', None),
                   getattr(section, 'even_page_footer', None)):
            if hf is None:
                continue
            for para in hf.paragraphs:
                for r_elem in para._element.findall('.//' + qn('w:r')):
                    _fix_run_text(r_elem)

    return total_changes


# =============================================================================
# 主格式化函数
# =============================================================================

def format_document(input_path, output_path=None, mode='full',
                    has_cover=False, has_toc=False,
                    page_num_style='dash',
                    header_left='', header_right='',
                    fix_punctuation=True,
                    landscape_wide_tables=False):
    """
    对Word文档执行自动排版。

    Args:
        input_path: 输入docx文件路径
        output_path: 输出docx文件路径（None则自动生成）
        mode: 'full' 完整排版 | 'light' 轻量排版
        has_cover: 是否有封面页
        has_toc: 是否有目录
        page_num_style: 'dash' ("— X —") | 'plain' (纯数字)
        header_left: 页眉左侧文本（报告名称）
        header_right: 页眉右侧文本（编制单位）
        landscape_wide_tables: 是否将宽表格强制切换为横向页。默认关闭，避免产生孤立横向页。

    Returns:
        output_path: 输出文件路径
        report: 格式化报告（字典）
    """
    if output_path is None:
        base, ext = os.path.splitext(input_path)
        output_path = f"{base}_formatted{ext}"

    doc = Document(input_path)

    # 重置目录书签计数器
    global _toc_bookmark_id
    _toc_bookmark_id = 0

    report = {
        'total_paragraphs': 0,
        'headings_formatted': 0,
        'body_formatted': 0,
        'tables_formatted': 0,
        'table_titles_formatted': 0,
        'figure_titles_formatted': 0,
        'appendix_formatted': 0,
        'punctuation_changes': 0,
        'warnings': [],
    }

    # ---- 0a. 重置文档默认样式的字体（清除主题字体） ----
    _reset_document_default_fonts(doc)

    # ---- 0b. 合并孤立序号段落 ----
    # 有些文档中序号（如"2." "（一）" "①"）和正文内容被分成两个段落，
    # 导致排版后序号独占一行。这里预处理合并它们。
    _merge_paragraphs(doc)

    # ---- 0c. 清除项目符号标记 ----
    # 清除段首的 "." "·" "•" 等项目符号字符和Word列表编号属性
    _clean_bullet_markers(doc)

    # ---- 0d. v2.7 新增：中文标点国标 GB/T 15834—2011 自动修正 ----
    # 修正 AI 生成中文文档最常见的半角假装全角问题：
    # 半角逗号/句号/冒号/分号/问号/叹号、ASCII 直引号、半角括号、尖括号冒充书名号、
    # -- 冒充破折号、... 冒充省略号、数值范围半角 -/~ 等。
    # 跳过白名单：URL、邮箱、版本号、千分位、小数点、行内代码。
    if fix_punctuation:
        punct_changes = fix_punctuation_in_doc(doc)
        if punct_changes > 0:
            print(f"  标点修正: 共修改 {punct_changes} 处中文标点")
            report['punctuation_changes'] = punct_changes
        else:
            report['punctuation_changes'] = 0

    # ---- 1. 页面设置 ----
    for section in doc.sections:
        setup_page(section)

    # ---- 2. 页眉 ----
    if header_left or header_right:
        setup_header(doc, header_left, header_right)

    # ---- 3. 页码 ----
    setup_page_numbers(doc, style=page_num_style, has_cover=has_cover)

    # ---- 4. 封面格式化 ----
    cover_range = None
    if has_cover:
        cover_range = detect_cover_paragraphs(doc)
        if not cover_range:
            # v2.7.3: 回退策略——原文档封面和正文间无分页符，
            # 按一级标题推断封面范围
            cover_range = _infer_cover_range(doc)
            if cover_range:
                print(f"  封面推断: 无分页符，推断封面范围 {cover_range}")
        if cover_range:
            format_cover_page(doc, cover_range)

    # ---- 4b. 封面结束后确保分页 ----
    # v2.7.3: 封面结束后必须强制另起页（目录或正文在新的一页开始）
    if cover_range:
        cover_end_para = doc.paragraphs[cover_range[1]]
        # 检查封面最后段落是否有分页符
        has_page_break = False
        for run in cover_end_para.runs:
            for child in run._element:
                if child.tag == qn('w:br'):
                    br_type = child.get(qn('w:type'))
                    if br_type == 'page':
                        has_page_break = True
                        break
        if not has_page_break:
            # 也检查是否有 page_break_before 属性
            # 找到封面后的第一个非空段落，给它设置 page_break_before
            for j in range(cover_range[1] + 1, len(doc.paragraphs)):
                next_para = doc.paragraphs[j]
                if next_para.text.strip():
                    next_para.paragraph_format.page_break_before = True
                    print(f"  封面分页: 段落{j} '{next_para.text.strip()[:20]}' 设置 page_break_before=True")
                    break

    # ---- 5. 逐段落格式化（跳过封面） ----
    in_toc = False
    toc_ended = False
    toc_title_idx = None   # v2.7.3: "目录"标题段落的索引
    headings_list = []     # v2.7.3: 收集所有标题，用于插入目录条目
    # headings_list 格式: [(level, text, bookmark_name), ...]
    _toc_bm_counter = 0   # 目录书签计数器

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        report['total_paragraphs'] += 1

        # 跳过封面段落（已单独处理）
        if cover_range and i <= cover_range[1]:
            continue

        # 跳过空段落
        if not text:
            continue

        # 检测目录区域
        if has_toc and not toc_ended:
            if text == '目录' or text == '目 录':
                in_toc = True
                toc_title_idx = i  # v2.7.3: 记录目录标题位置
                # 格式化目录标题
                set_paragraph_font(para, TOC_TITLE_FONT, FONT_TNR,
                                   TOC_TITLE_SIZE, TOC_TITLE_BOLD)
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.page_break_before = True
                continue
            if in_toc:
                # 目录区域在遇到一级标题时结束
                if RE_H1.match(text):
                    in_toc = False
                    toc_ended = True
                else:
                    # 目录页内容不做自动格式化（由Word目录功能控制）
                    continue

        # 检测标题层级（传入 para 以便回退到 Word 样式检测）
        level = detect_heading_level(text, para)

        if level is not None:
            format_heading(para, level)
            report['headings_formatted'] += 1

            # v2.7.3: 收集标题文本（用于目录条目生成），并添加书签
            if level <= 2:
                _toc_bm_counter += 1
                bm_name = f'_Toc_{_toc_bm_counter}'
                add_toc_bookmark(para, bm_name)
                headings_list.append((level, text, bm_name))

            # v2.7.3: 分页规则——只在以下位置强制分页：
            # 1. 每个一级标题另起页（包括第一个，因为前面是封面或目录）
            # 2. 每个附件标题另起页
            # 3. 二级及以下标题、正文、图表标题等不强制分页
            if level == 1:
                para.paragraph_format.page_break_before = True
            else:
                para.paragraph_format.page_break_before = False

            # 检查：页面最后一行不应出现标题
            # （这个需要在排版后人工检查，脚本只能设置keep_with_next）

        elif is_table_title(text):
            format_table_title_paragraph(para)
            report['table_titles_formatted'] += 1
            # v2.7.3: 表格标题不强制分页
            para.paragraph_format.page_break_before = False

        elif is_figure_title(text):
            format_figure_title_paragraph(para)
            report['figure_titles_formatted'] += 1
            # v2.7.3: 图标题不强制分页
            para.paragraph_format.page_break_before = False

        elif RE_APPENDIX.match(text):
            # v2.7.3: 附件标题另起页，格式同一级标题（黑体加粗）
            set_paragraph_font(para, H1_FONT, FONT_TNR, H1_SIZE, H1_BOLD,
                               color=RGBColor(0, 0, 0))
            set_paragraph_format(
                para,
                first_indent=Cm(0),  # 附件顶格排版，不缩进
                line_spacing=BODY_LINE_SPACING,
                space_before=BODY_PARA_BEFORE,
                space_after=BODY_PARA_AFTER,
                keep_with_next=True,
            )
            set_outline_level(para, 0)  # 附件标题也用大纲级别0
            para.paragraph_format.page_break_before = True
            # v2.7.3: 附件标题加入目录条目列表（level=1，与H1同级）
            _toc_bm_counter += 1
            bm_name = f'_Toc_{_toc_bm_counter}'
            add_toc_bookmark(para, bm_name)
            headings_list.append((1, text, bm_name))
            report['appendix_formatted'] += 1

        else:
            # 正文段落
            format_body_paragraph(para)
            report['body_formatted'] += 1
            # v2.7.3: 正文不强制分页
            para.paragraph_format.page_break_before = False

    # ---- 5b. v2.7.3: 插入目录条目 ----
    if has_toc and headings_list:
        if toc_title_idx is None:
            # v2.7.3: 原文档没有"目录"段落，自动在封面后插入目录页
            # 确定插入位置：封面结束后，第一个H1之前
            body = doc.element.body
            all_paras = list(body.findall(qn('w:p')))
            insert_para = None
            if cover_range:
                insert_para = all_paras[cover_range[1]] if cover_range[1] < len(all_paras) else None
            else:
                # 没有封面，在文档开头插入（第一个段落之前）
                insert_para = all_paras[0] if all_paras else None

            if insert_para is not None:
                # 创建"目录"标题段落：黑体小二号加粗，居中
                # v2.7.5: 添加 pageBreakBefore，确保目录页另起新页（在封面后）
                toc_title_xml = (
                    f'<w:p {nsdecls("w")}>'
                    f'  <w:pPr>'
                    f'    <w:pageBreakBefore/>'
                    f'    <w:jc w:val="center"/>'
                    f'    <w:spacing w:before="0" w:after="240" w:line="360" w:lineRule="auto"/>'
                    f'    <w:rPr>'
                    f'      <w:rFonts w:eastAsia="{FONT_HEITI}" w:ascii="{FONT_TNR}" w:hAnsi="{FONT_TNR}"/>'
                    f'      <w:b/>'
                    f'      <w:sz w:val="44"/>'
                    f'      <w:szCs w:val="44"/>'
                    f'      <w:color w:val="000000"/>'
                    f'    </w:rPr>'
                    f'  </w:pPr>'
                    f'  <w:r>'
                    f'    <w:rPr>'
                    f'      <w:rFonts w:eastAsia="{FONT_HEITI}" w:ascii="{FONT_TNR}" w:hAnsi="{FONT_TNR}"/>'
                    f'      <w:b/>'
                    f'      <w:sz w:val="44"/>'
                    f'      <w:szCs w:val="44"/>'
                    f'      <w:color w:val="000000"/>'
                    f'    </w:rPr>'
                    f'    <w:t>目录</w:t>'
                    f'  </w:r>'
                    f'</w:p>'
                )
                toc_title_elem = parse_xml(toc_title_xml)

                if cover_range:
                    # 在封面最后一个段落后插入
                    insert_para.addnext(toc_title_elem)
                else:
                    # 在文档第一个段落前插入
                    insert_para.addprevious(toc_title_elem)

                # 重新计算 toc_title_idx
                all_paras = list(body.findall(qn('w:p')))
                for i, p_elem in enumerate(all_paras):
                    texts = p_elem.findall('.//' + qn('w:t'))
                    text = ''.join(t.text or '' for t in texts).strip()
                    if text == '目录':
                        toc_title_idx = i
                        break

                print(f"  目录页: 自动插入于封面后（段落 {toc_title_idx}）")

        if toc_title_idx is not None:
            insert_toc_entries(doc, toc_title_idx, headings_list,
                               cover_end_idx=cover_range[1] if cover_range else None)
            print(f"  目录条目: 插入 {len([h for h in headings_list if h[0] <= 2])} 条（含独立书签）")

    # ---- 6. 表格格式化 ----
    wide_table_indices = []
    for idx, table in enumerate(doc.tables):
        format_table(table)
        report['tables_formatted'] += 1
        # 检测是否需要横向排版
        if _estimate_table_width_need(table) == 'wide':
            wide_table_indices.append(idx)

    # ---- 6b. 按需为宽表格插入横向分节 ----
    if wide_table_indices and landscape_wide_tables:
        _wrap_wide_tables_landscape(doc, wide_table_indices)
        print(f"  宽表格横向排版: {len(wide_table_indices)} 个")
    elif wide_table_indices:
        report['warnings'].append(
            f"检测到 {len(wide_table_indices)} 个宽表格，已保留纵向页面并自动适应宽度；"
            "如确需横向页，请显式添加 --landscape-wide-tables"
        )

    # ---- 7. 保存 ----
    doc.save(output_path)

    return output_path, report


# =============================================================================
# CLI 入口
# =============================================================================

def main():
    parser = argparse.ArgumentParser(
        description='博维咨询Word报告自动排版工具（依据V2.6标准）'
    )
    parser.add_argument('input', help='输入docx文件路径')
    parser.add_argument('output', nargs='?', default=None, help='输出docx文件路径')
    parser.add_argument('--mode', choices=['full', 'light'], default='full',
                        help='排版模式: full=完整排版, light=轻量排版')
    parser.add_argument('--cover', action='store_true', help='第一页为封面页')
    parser.add_argument('--toc', action='store_true', help='文档包含目录')
    parser.add_argument('--page-num-style', choices=['dash', 'plain'], default='dash',
                        help='页码样式: dash="— X —", plain=纯数字')
    parser.add_argument('--header-left', default='', help='页眉左侧文本（报告名称）')
    parser.add_argument('--header-right', default='', help='页眉右侧文本（编制单位）')
    parser.add_argument('--no-fix-punctuation', dest='fix_punctuation',
                        action='store_false', default=True,
                        help='关闭 GB/T 15834 中文标点自动修正（默认开启）')
    parser.add_argument('--landscape-wide-tables', action='store_true',
                        help='将检测到的宽表格强制切换为横向页（默认关闭）')

    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"错误: 文件不存在 - {args.input}")
        sys.exit(1)

    print(f"正在排版: {args.input}")
    print(f"排版模式: {args.mode}")

    output_path, report = format_document(
        args.input,
        args.output,
        mode=args.mode,
        has_cover=args.cover,
        has_toc=args.toc,
        page_num_style=args.page_num_style,
        header_left=args.header_left,
        header_right=args.header_right,
        fix_punctuation=args.fix_punctuation,
        landscape_wide_tables=args.landscape_wide_tables,
    )

    print(f"\n排版完成! 输出文件: {output_path}")
    print(f"  段落总数: {report['total_paragraphs']}")
    print(f"  标题格式化: {report['headings_formatted']}")
    print(f"  正文格式化: {report['body_formatted']}")
    print(f"  表格格式化: {report['tables_formatted']}")
    print(f"  表标题格式化: {report['table_titles_formatted']}")
    print(f"  图标题格式化: {report['figure_titles_formatted']}")
    if report.get('punctuation_changes', 0) > 0:
        print(f"  中文标点修正: {report['punctuation_changes']} 处")

    if report['warnings']:
        print(f"\n⚠ 警告:")
        for w in report['warnings']:
            print(f"  - {w}")


if __name__ == '__main__':
    main()
