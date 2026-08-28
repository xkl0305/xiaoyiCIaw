#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word 文档解析模块
支持 .docx / .doc / .md 解析
"""

import os
import re
from pathlib import Path

def parse_docx(file_path):
    """
    解析 .docx 文件
    
    Args:
        file_path: .docx 文件路径
    
    Returns:
        dict: 文档结构
    """
    result = {
        "success": False,
        "paragraphs": [],
        "tables": [],
        "images": [],
        "headers": [],
        "footers": [],
        "error": None
    }
    
    try:
        from docx import Document
        from docx.oxml import parse_xml
        
        doc = Document(file_path)
        
        # 1. 解析段落
        for para in doc.paragraphs:
            para_info = {
                "text": para.text,
                "style": para.style.name if para.style else "Normal",
                "alignment": get_alignment(para),
                "level": get_heading_level(para),
                "runs": []
            }
            
            # 解析 runs（文本片段）
            for run in para.runs:
                run_info = {
                    "text": run.text,
                    "bold": run.bold,
                    "italic": run.italic,
                    "font_name": run.font.name,
                    "font_size": run.font.size.pt if run.font.size else None,
                    "color": str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None
                }
                para_info["runs"].append(run_info)
            
            result["paragraphs"].append(para_info)
        
        # 2. 解析表格
        for table in doc.tables:
            table_info = {
                "rows": len(table.rows),
                "cols": len(table.columns),
                "cells": []
            }
            
            for row in table.rows:
                for cell in row.cells:
                    cell_info = {
                        "text": cell.text,
                        "rowspan": 1,  # 简化
                        "colspan": 1   # 简化
                    }
                    table_info["cells"].append(cell_info)
            
            result["tables"].append(table_info)
        
        # 3. 解析图片（简化）
        for rel in doc.part.rels:
            if "image" in doc.part.rels[rel].reltype:
                result["images"].append({
                    "id": rel,
                    "type": "image"
                })
        
        # 4. 解析页眉页脚（简化）
        for section in doc.sections:
            if section.header:
                result["headers"].append({
                    "text": section.header.paragraphs[0].text if section.header.paragraphs else ""
                })
            if section.footer:
                result["footers"].append({
                    "text": section.footer.paragraphs[0].text if section.footer.paragraphs else ""
                })
        
        result["success"] = True
        
    except Exception as e:
        result["error"] = str(e)
    
    return result

def parse_doc(file_path):
    """
    解析 .doc 文件（需先转换为 .docx）
    
    Args:
        file_path: .doc 文件路径
    
    Returns:
        dict: 文档结构
    """
    result = {
        "success": False,
        "paragraphs": [],
        "tables": [],
        "images": [],
        "error": None,
        "warning": "请先将 .doc 文件另存为 .docx 格式"
    }
    
    try:
        # 尝试使用 LibreOffice 转换
        import subprocess
        
        output_dir = os.path.dirname(file_path)
        cmd = [
            "soffice",
            "--headless",
            "--convert-to", "docx",
            "--outdir", output_dir,
            file_path
        ]
        
        subprocess.run(cmd, check=True, capture_output=True)
        
        # 转换成功后解析 .docx
        docx_path = os.path.splitext(file_path)[0] + ".docx"
        if os.path.exists(docx_path):
            result = parse_docx(docx_path)
            result["warning"] = f".doc 已自动转换为 .docx：{docx_path}"
        
    except Exception as e:
        result["error"] = f".doc 解析失败，请手动转换为 .docx：{str(e)}"
    
    return result

def parse_md(file_path):
    """
    解析 .md 文件
    
    Args:
        file_path: .md 文件路径
    
    Returns:
        dict: 文档结构
    """
    result = {
        "success": False,
        "paragraphs": [],
        "tables": [],
        "images": [],
        "error": None
    }
    
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        
        # 简单解析 Markdown
        lines = content.split("\n")
        
        for i, line in enumerate(lines):
            para_info = {
                "text": line,
                "style": "Normal",
                "alignment": "left",
                "level": None,
                "runs": [{"text": line, "bold": False, "italic": False}]
            }
            
            # 识别标题
            if line.startswith("# "):
                para_info["level"] = 1
                para_info["style"] = "Heading 1"
            elif line.startswith("## "):
                para_info["level"] = 2
                para_info["style"] = "Heading 2"
            elif line.startswith("### "):
                para_info["level"] = 3
                para_info["style"] = "Heading 3"
            
            result["paragraphs"].append(para_info)
        
        result["success"] = True
        
    except Exception as e:
        result["error"] = str(e)
    
    return result

def get_alignment(paragraph):
    """获取段落对齐方式"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    alignment_map = {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify"
    }
    
    if paragraph.alignment:
        return alignment_map.get(paragraph.alignment, "left")
    return "left"

def get_heading_level(paragraph):
    """获取标题层级"""
    style_name = paragraph.style.name if paragraph.style else ""
    
    # 识别标题样式
    if style_name.startswith("Heading"):
        try:
            level = int(style_name.split()[-1])
            return level
        except:
            pass
    
    # 正则识别中文标题
    text = paragraph.text.strip()
    
    # 一级标题：一、二、三、
    if re.match(r"^[一二三四五六七八九十]+、", text):
        return 1
    
    # 二级标题：（一）（二）（三）
    if re.match(r"^（[一二三四五六七八九十]+）", text):
        return 2
    
    # 三级标题：1. 2. 3.
    if re.match(r"^[0-9]+\.", text):
        return 3
    
    # 四级标题：（1）（2）（3）
    if re.match(r"^（[0-9]+）", text):
        return 4
    
    return None

def identify_gov_elements(paragraphs):
    """
    识别公文要素
    
    Args:
        paragraphs: 段落列表
    
    Returns:
        dict: 公文要素
    """
    gov_elements = {
        "title": None,
        "main_sender": None,
        "body_start": None,
        "attachment": None,
        "signature": None,
        "date": None,
        "cc": None,
        "issuer": None
    }
    
    for i, para in enumerate(paragraphs):
        text = para["text"].strip()
        
        # 标题（关于…的通知/报告/请示等）
        if re.search(r"关于.+的(通知|报告|请示|批复|函|意见|通报|公告|决定)", text):
            gov_elements["title"] = text
        
        # 主送机关（冒号结尾）
        if text.endswith("：") and len(text) < 50:
            gov_elements["main_sender"] = text
        
        # 落款机关
        if re.search(r"[省市区县乡镇]+(人民)?(政府|厅|局|委员会|办公室)$", text):
            gov_elements["signature"] = text
        
        # 日期
        if re.search(r"[0-9]{4}年[0-9]+月[0-9]+日", text):
            gov_elements["date"] = text
        
        # 抄送
        if text.startswith("抄送："):
            gov_elements["cc"] = text
        
        # 附件
        if text.startswith("附件："):
            gov_elements["attachment"] = text
    
    return gov_elements

if __name__ == "__main__":
    # 测试
    import sys
    
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        
        if file_path.endswith(".docx"):
            result = parse_docx(file_path)
        elif file_path.endswith(".doc"):
            result = parse_doc(file_path)
        elif file_path.endswith(".md"):
            result = parse_md(file_path)
        else:
            print(f"不支持的文件格式：{file_path}")
            sys.exit(1)
        
        import json
        print(json.dumps(result, ensure_ascii=False, indent=2))
