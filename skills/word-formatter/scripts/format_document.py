#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word 智能排版主脚本
支持：文档解析、模板应用、公文国标排版、差异报告生成
"""

import sys
import json
import os
import re
from pathlib import Path
from datetime import datetime
import shutil

# 导入同级模块
from parse_docx import parse_docx, parse_doc
from template_engine import load_template, apply_template
from gov_formatter import format_gov_document, check_gov_standard
from generate_diff import generate_diff_report

def format_document(source_file, template_id="business_default", gov_doctype=None, 
                    format_options=None, output_dir=None):
    """
    主排版函数
    
    Args:
        source_file: 源 Word 文档路径
        template_id: 模板 ID
        gov_doctype: 公文类型（仅 gov_ 模板）
        format_options: 格式选项 dict
        output_dir: 输出目录
    
    Returns:
        dict: 包含输出文件路径、报告路径等
    """
    result = {
        "success": False,
        "formatted_file": None,
        "diff_report": None,
        "gov_check_report": None,
        "summary": "",
        "md_archive_path": None,
        "errors": []
    }
    
    try:
        # 1. 验证输入文件
        if not os.path.exists(source_file):
            result["errors"].append(f"文件不存在：{source_file}")
            return result
        
        # 2. 解析文档
        print(f"📂 正在解析文档：{source_file}")
        doc_structure = parse_docx(source_file) if source_file.endswith('.docx') else parse_doc(source_file)
        
        if not doc_structure["success"]:
            result["errors"].append(f"文档解析失败：{doc_structure.get('error', '未知错误')}")
            return result
        
        # 3. 加载模板
        print(f"🎨 正在加载模板：{template_id}")
        template = load_template(template_id, gov_doctype)
        
        if not template["success"]:
            result["errors"].append(f"模板加载失败：{template.get('error', '未知错误')}")
            # 回退到默认模板
            print("⚠️ 回退到默认模板：simple_minimal")
            template = load_template("simple_minimal")
        
        # 4. 应用排版规则
        print(f"🔧 正在应用排版规则...")
        format_options = format_options or {}
        formatted_doc = apply_template(doc_structure, template, format_options)
        
        # 5. 公文专属处理
        is_gov_template = template_id.startswith("gov_")
        gov_check_report = None
        
        if is_gov_template:
            print(f"🏛️ 正在应用公文国标排版...")
            formatted_doc = format_gov_document(formatted_doc, template, gov_doctype)
            
            # 公文国标校验
            if format_options.get("gov_check", True):
                print(f"✅ 正在生成公文国标校验报告...")
                gov_check_report = check_gov_standard(formatted_doc, template)
                result["gov_check_report"] = gov_check_report["report_path"]
        
        # 6. 生成输出文件
        if output_dir is None:
            output_dir = os.path.join(os.path.dirname(__file__), "../outputs", datetime.now().strftime("%Y-%m-%d"))
            os.makedirs(output_dir, exist_ok=True)
        
        output_filename = os.path.splitext(os.path.basename(source_file))[0] + "_排版后.docx"
        output_path = os.path.join(output_dir, output_filename)
        
        print(f"💾 正在保存排版后文档：{output_path}")
        save_formatted_doc(formatted_doc, output_path)
        result["formatted_file"] = output_path
        
        # 7. 生成差异对比报告
        if format_options.get("generate_diff_report", True):
            print(f"📊 正在生成差异对比报告...")
            diff_report_path = generate_diff_report(source_file, output_path, doc_structure, formatted_doc)
            result["diff_report"] = diff_report_path
        
        # 8. Markdown 归档
        print(f"📝 正在归档 Markdown...")
        md_archive_path = archive_to_markdown(source_file, output_path, doc_structure, formatted_doc)
        result["md_archive_path"] = md_archive_path
        
        # 9. 生成摘要
        result["summary"] = generate_summary(doc_structure, template, is_gov_template, gov_check_report)
        result["success"] = True
        
        print(f"✅ 排版完成！")
        
    except Exception as e:
        result["errors"].append(f"排版失败：{str(e)}")
        import traceback
        traceback.print_exc()
    
    return result

def save_formatted_doc(doc, output_path):
    """保存排版后的文档"""
    from docx import Document
    
    if isinstance(doc, dict):
        # 从结构重建文档
        document = Document()
        
        # 应用页面设置
        if "page_setup" in doc:
            apply_page_setup(document, doc["page_setup"])
        
        # 添加段落
        for para_info in doc.get("paragraphs", []):
            para = document.add_paragraph()
            apply_paragraph_format(para, para_info)
            
            # 添加文本
            for run_info in para_info.get("runs", []):
                run = para.add_run(run_info["text"])
                apply_run_format(run, run_info)
        
        # 添加表格
        for table_info in doc.get("tables", []):
            table = document.add_table(rows=table_info["rows"], cols=table_info["cols"])
            apply_table_format(table, table_info)
        
        document.save(output_path)
    else:
        # 已经是 Document 对象
        doc.save(output_path)

def apply_page_setup(document, page_setup):
    """应用页面设置"""
    from docx.shared import Cm
    
    section = document.sections[0]
    
    # 页边距
    if "margin" in page_setup:
        margin = page_setup["margin"]
        section.top_margin = Cm(margin.get("top", 2.54))
        section.bottom_margin = Cm(margin.get("bottom", 2.54))
        section.left_margin = Cm(margin.get("left", 2.54))
        section.right_margin = Cm(margin.get("right", 2.54))
    
    # 纸张大小
    if "paper" in page_setup:
        from docx.shared import Mm
        if page_setup["paper"] == "A4":
            section.page_width = Mm(210)
            section.page_height = Mm(297)

def apply_paragraph_format(para, para_info):
    """应用段落格式"""
    from docx.shared import Pt, Cm
    
    fmt = para.paragraph_format
    
    # 对齐方式
    if "alignment" in para_info:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        fmt.alignment = align_map.get(para_info["alignment"], WD_ALIGN_PARAGRAPH.LEFT)
    
    # 行距
    if "line_spacing" in para_info:
        fmt.line_spacing = Pt(para_info["line_spacing"])
    
    # 段间距
    if "space_before" in para_info:
        fmt.space_before = Pt(para_info["space_before"])
    if "space_after" in para_info:
        fmt.space_after = Pt(para_info["space_after"])
    
    # 首行缩进
    if "first_line_indent" in para_info:
        fmt.first_line_indent = Cm(para_info["first_line_indent"] * 0.74)  # 字符转厘米

def apply_run_format(run, run_info):
    """应用文本格式"""
    from docx.shared import Pt, RGBColor
    
    # 字体
    if "font_name" in run_info:
        run.font.name = run_info["font_name"]
    
    # 字号
    if "font_size" in run_info:
        run.font.size = Pt(run_info["font_size"])
    
    # 加粗
    if "bold" in run_info:
        run.font.bold = run_info["bold"]
    
    # 斜体
    if "italic" in run_info:
        run.font.italic = run_info["italic"]
    
    # 颜色
    if "color" in run_info:
        run.font.color.rgb = RGBColor.from_string(run_info["color"])

def apply_table_format(table, table_info):
    """应用表格格式"""
    # 简化实现
    pass

def archive_to_markdown(source_file, formatted_file, doc_structure, formatted_doc):
    """归档到 Markdown"""
    archive_dir = os.path.join(os.path.dirname(__file__), "../history", datetime.now().strftime("%Y-%m-%d"))
    os.makedirs(archive_dir, exist_ok=True)
    
    # 归档原始文档结构
    source_md_path = os.path.join(archive_dir, os.path.splitext(os.path.basename(source_file))[0] + "_排版前.md")
    with open(source_md_path, "w", encoding="utf-8") as f:
        f.write(f"# 排版前文档结构\n\n")
        f.write(f"## 文档信息\n")
        f.write(f"- 文件名：{os.path.basename(source_file)}\n")
        f.write(f"- 段落数：{len(doc_structure.get('paragraphs', []))}\n")
        f.write(f"- 表格数：{len(doc_structure.get('tables', []))}\n")
        f.write(f"- 图片数：{len(doc_structure.get('images', []))}\n\n")
        
        f.write(f"## 段落内容\n\n")
        for i, para in enumerate(doc_structure.get("paragraphs", [])[:20]):  # 只显示前 20 段
            f.write(f"### 段落 {i+1}\n")
            f.write(f"```\n{para.get('text', '')}\n```\n\n")
    
    # 归档排版后文档结构
    formatted_md_path = os.path.join(archive_dir, os.path.splitext(os.path.basename(source_file))[0] + "_排版后.md")
    with open(formatted_md_path, "w", encoding="utf-8") as f:
        f.write(f"# 排版后文档结构\n\n")
        f.write(f"## 文档信息\n")
        f.write(f"- 文件名：{os.path.basename(formatted_file)}\n")
        f.write(f"- 应用模板：{doc_structure.get('template_id', 'unknown')}\n\n")
        f.write(f"## 排版规则已应用\n")
        f.write(f"- 字体规范化\n")
        f.write(f"- 字号规范化\n")
        f.write(f"- 行距调整\n")
        f.write(f"- 段落间距调整\n")
        f.write(f"- 标题样式应用\n")
    
    return archive_dir

def generate_summary(doc_structure, template, is_gov_template, gov_check_report):
    """生成排版摘要"""
    summary = f"✅ 排版完成\n\n"
    summary += f"- 识别段落：{len(doc_structure.get('paragraphs', []))} 个\n"
    summary += f"- 识别表格：{len(doc_structure.get('tables', []))} 个\n"
    summary += f"- 识别图片：{len(doc_structure.get('images', []))} 张\n"
    summary += f"- 应用模板：{template.get('name', template.get('template_id', 'unknown'))}\n"
    
    if is_gov_template:
        summary += f"\n🏛️ 公文排版已应用\n"
        if gov_check_report:
            passed = gov_check_report.get("passed", 0)
            total = gov_check_report.get("total", 0)
            summary += f"- 国标校验：{passed}/{total} 通过\n"
    
    return summary

def main():
    """命令行入口"""
    if len(sys.argv) < 2:
        print("用法：python format_document.py <source_file> [template_id] [options]")
        print("示例：python format_document.py 项目报告.docx biz_report")
        sys.exit(1)
    
    source_file = sys.argv[1]
    template_id = sys.argv[2] if len(sys.argv) > 2 else "business_default"
    
    # 解析选项
    format_options = {}
    for arg in sys.argv[3:]:
        if arg.startswith("--"):
            key, value = arg[2:].split("=")
            format_options[key] = value.lower() == "true"
    
    result = format_document(source_file, template_id, format_options=format_options)
    
    print("\n" + "="*50)
    print(json.dumps(result, ensure_ascii=False, indent=2))
    print("="*50)

if __name__ == "__main__":
    main()
